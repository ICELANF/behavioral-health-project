"""
Convert CLAUDE_INTEGRATION.md to Word format
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_word_document():
    doc = Document()

    # Set document title
    title = doc.add_heading('Claude API 与行为健康平台集成方案', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add metadata
    meta = doc.add_paragraph()
    meta.add_run('生成时间: 2026-01-26\n').italic = True
    meta.add_run('基于: Claude Platform 官方文档').italic = True
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacing

    # ===== Section 1 =====
    doc.add_heading('一、集成架构总览', level=1)

    doc.add_heading('八爪鱼架构中的定位', level=2)
    p = doc.add_paragraph()
    p.add_run('在八爪鱼架构中，Claude API 作为云端大脑，与本地 Ollama 形成双引擎模式：')

    # Architecture diagram as text
    arch_text = """
                    ┌─────────────────────────────────┐
                    │           大脑 (Brain)           │
                    ├─────────────────────────────────┤
                    │   Ollama (本地)  │  Claude API   │
                    │   qwen2.5/deepseek │  (云端)     │
                    └─────────────────────────────────┘
                                    │
            ┌───────────────────────┼───────────────────────┐
            ▼                       ▼                       ▼
       触手1:专家            触手2:用户行为          触手3:教练培养
       Chatflow                 养成                    体系
    """
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(arch_text)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)

    doc.add_heading('双引擎模式', level=2)

    # Table: dual engine
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '引擎'
    hdr_cells[1].text = '场景'
    hdr_cells[2].text = '优势'
    hdr_cells[3].text = '成本'

    row1 = table.rows[1].cells
    row1[0].text = 'Ollama (本地)'
    row1[1].text = '日常对话、简单问答'
    row1[2].text = '免费、隐私、低延迟'
    row1[3].text = '硬件成本'

    row2 = table.rows[2].cells
    row2[0].text = 'Claude API (云端)'
    row2[1].text = '复杂推理、工具调用、深度分析'
    row2[2].text = '高质量、强能力'
    row2[3].text = 'API费用'

    doc.add_paragraph()

    # ===== Section 2 =====
    doc.add_heading('二、直接集成方案', level=1)

    doc.add_heading('2.1 作为替代大脑 (Alternative Brain)', level=2)
    doc.add_paragraph('场景: 当本地 Ollama 模型能力不足时，切换到 Claude API')
    doc.add_paragraph('适用场景:')
    doc.add_paragraph('• 复杂健康方案制定', style='List Bullet')
    doc.add_paragraph('• 多步骤推理任务', style='List Bullet')
    doc.add_paragraph('• 需要工具调用的场景', style='List Bullet')

    doc.add_heading('2.2 Tool Use 集成 (工具调用)', level=2)
    doc.add_paragraph('Claude API 的 Tool Use 功能可以让 AI 调用我们定义的健康工具。')

    # Tool table
    doc.add_paragraph('可定义的健康工具清单:')
    tool_table = doc.add_table(rows=8, cols=3)
    tool_table.style = 'Table Grid'
    tool_hdr = tool_table.rows[0].cells
    tool_hdr[0].text = '工具名称'
    tool_hdr[1].text = '功能'
    tool_hdr[2].text = '所属触手'

    tools_data = [
        ('ttm_assessment', 'TTM阶段评估', '触手2'),
        ('diet_analyzer', '饮食分析', '触手1/2'),
        ('exercise_tracker', '运动记录查询', '触手2'),
        ('goal_setter', 'SMART目标设定', '触手2'),
        ('coach_evaluator', '教练能力评估', '触手3'),
        ('case_analyzer', '案例分析', '触手3'),
        ('knowledge_search', '知识库检索', '全部'),
    ]
    for i, (name, func, tentacle) in enumerate(tools_data, 1):
        row = tool_table.rows[i].cells
        row[0].text = name
        row[1].text = func
        row[2].text = tentacle

    doc.add_paragraph()

    doc.add_heading('2.3 Agent Skills 集成', level=2)
    doc.add_paragraph('Claude 的 Agent Skills 可以创建专业化的健康教练技能，例如动机访谈技能。')

    # ===== Section 3 =====
    doc.add_heading('三、间接集成方案', level=1)

    doc.add_heading('3.1 Dify + Claude API 混合架构', level=2)
    doc.add_paragraph('配置步骤:')
    doc.add_paragraph('1. 在 Dify 后台添加 Claude 模型供应商', style='List Number')
    doc.add_paragraph('2. 输入 Claude API Key', style='List Number')
    doc.add_paragraph('3. 在 Chatflow 中配置模型切换逻辑', style='List Number')

    doc.add_heading('3.2 FastAPI 后端集成', level=2)
    doc.add_paragraph('在后端 API 层集成 Claude，实现业务逻辑。可以创建 AIService 类同时管理 Ollama 和 Claude 客户端。')

    doc.add_heading('3.3 知识库增强 (RAG + Claude)', level=2)
    doc.add_paragraph('利用 Claude 的 100万 token 上下文窗口，增强知识检索能力，实现深度理解和整合。')

    # ===== Section 4 =====
    doc.add_heading('四、特性对接矩阵', level=1)

    feature_table = doc.add_table(rows=9, cols=4)
    feature_table.style = 'Table Grid'
    f_hdr = feature_table.rows[0].cells
    f_hdr[0].text = 'Claude 特性'
    f_hdr[1].text = '平台应用'
    f_hdr[2].text = '触手'
    f_hdr[3].text = '优先级'

    features = [
        ('Tool Use', '健康工具调用', '全部', 'P0'),
        ('Extended Thinking', '复杂方案制定', '触手1/2', 'P1'),
        ('Citations', '知识来源标注', '触手1', 'P1'),
        ('Web Search', '最新健康资讯', '触手1', 'P2'),
        ('Code Execution', '数据分析可视化', '触手3', 'P2'),
        ('PDF Support', '病历/报告解读', '触手2', 'P1'),
        ('Prompt Caching', '降低成本', '全部', 'P1'),
        ('Batch API', '批量评估', '触手3', 'P2'),
    ]
    for i, (feat, app, tent, pri) in enumerate(features, 1):
        row = feature_table.rows[i].cells
        row[0].text = feat
        row[1].text = app
        row[2].text = tent
        row[3].text = pri

    doc.add_paragraph()

    # ===== Section 5 =====
    doc.add_heading('五、实施路线图', level=1)

    doc.add_heading('第一阶段: 基础集成', level=2)
    doc.add_paragraph('• 在 Dify 中配置 Claude 模型供应商', style='List Bullet')
    doc.add_paragraph('• 创建 Claude 版本的专家 Chatflow', style='List Bullet')
    doc.add_paragraph('• 实现 Ollama/Claude 智能路由', style='List Bullet')
    doc.add_paragraph('• 测试基本对话功能', style='List Bullet')

    doc.add_heading('第二阶段: Tool Use 集成', level=2)
    doc.add_paragraph('• 定义健康工具 JSON Schema', style='List Bullet')
    doc.add_paragraph('• 实现 ttm_assessment、diet_analyzer、goal_setter 工具', style='List Bullet')
    doc.add_paragraph('• 在 FastAPI 中创建工具执行端点', style='List Bullet')
    doc.add_paragraph('• 测试工具调用完整流程', style='List Bullet')

    doc.add_heading('第三阶段: 高级特性', level=2)
    doc.add_paragraph('• 集成 Extended Thinking (复杂方案)', style='List Bullet')
    doc.add_paragraph('• 集成 Citations (知识引用)', style='List Bullet')
    doc.add_paragraph('• 配置 Prompt Caching (成本优化)', style='List Bullet')
    doc.add_paragraph('• 集成 PDF Support (报告解读)', style='List Bullet')

    # ===== Section 6 =====
    doc.add_heading('六、成本估算', level=1)

    doc.add_heading('Claude API 定价 (参考)', level=2)
    price_table = doc.add_table(rows=4, cols=3)
    price_table.style = 'Table Grid'
    p_hdr = price_table.rows[0].cells
    p_hdr[0].text = '模型'
    p_hdr[1].text = '输入 (每百万token)'
    p_hdr[2].text = '输出 (每百万token)'

    prices = [
        ('Claude Sonnet 4', '$3', '$15'),
        ('Claude Opus 4', '$15', '$75'),
        ('Claude Haiku 3.5', '$0.8', '$4'),
    ]
    for i, (model, inp, out) in enumerate(prices, 1):
        row = price_table.rows[i].cells
        row[0].text = model
        row[1].text = inp
        row[2].text = out

    doc.add_paragraph()
    doc.add_heading('成本优化策略', level=2)
    doc.add_paragraph('• 简单对话 → Ollama (免费)', style='List Bullet')
    doc.add_paragraph('• 复杂推理 → Claude Sonnet (性价比)', style='List Bullet')
    doc.add_paragraph('• 深度分析 → Claude Opus (最强能力)', style='List Bullet')
    doc.add_paragraph('• 批量处理 → Batch API (50%折扣)', style='List Bullet')
    doc.add_paragraph('• 重复提示 → Prompt Caching (90%折扣)', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('预估月成本: $50-200 (中等使用量)')

    # ===== Section 7 =====
    doc.add_heading('七、环境配置', level=1)

    doc.add_heading('新增环境变量', level=2)
    env_code = """# .env 新增

# Claude API 配置
CLAUDE_API_KEY=sk-ant-api03-xxxxxx
CLAUDE_MODEL_DEFAULT=claude-sonnet-4-20250514
CLAUDE_MODEL_ADVANCED=claude-opus-4-20250514
CLAUDE_MODEL_FAST=claude-3-5-haiku-20241022

# 智能路由配置
AI_ROUTER_ENABLED=true
AI_ROUTER_COMPLEXITY_THRESHOLD=0.7"""

    code_p = doc.add_paragraph()
    code_r = code_p.add_run(env_code)
    code_r.font.name = 'Consolas'
    code_r.font.size = Pt(10)

    doc.add_heading('依赖安装', level=2)
    doc.add_paragraph('pip install anthropic>=0.40.0')

    # ===== Section 8 =====
    doc.add_heading('八、代码示例', level=1)
    doc.add_paragraph('完整的 Tool Use 示例代码请参考 Markdown 原文档或项目源码。')
    doc.add_paragraph('核心流程:')
    doc.add_paragraph('1. 定义健康工具的 JSON Schema', style='List Number')
    doc.add_paragraph('2. 创建 ClaudeToolService 类', style='List Number')
    doc.add_paragraph('3. 实现 process_tool_call 方法处理工具调用', style='List Number')
    doc.add_paragraph('4. 在对话循环中处理 tool_use 响应', style='List Number')

    # ===== Section 9 =====
    doc.add_heading('九、总结', level=1)

    doc.add_heading('集成价值', level=2)
    value_table = doc.add_table(rows=5, cols=2)
    value_table.style = 'Table Grid'
    v_hdr = value_table.rows[0].cells
    v_hdr[0].text = '价值点'
    v_hdr[1].text = '说明'

    values = [
        ('能力增强', 'Claude 补充 Ollama 的复杂推理能力'),
        ('工具调用', '实现健康工具的智能调用'),
        ('成本可控', '智能路由 + 缓存优化成本'),
        ('可扩展性', '双引擎架构便于未来扩展'),
    ]
    for i, (point, desc) in enumerate(values, 1):
        row = value_table.rows[i].cells
        row[0].text = point
        row[1].text = desc

    doc.add_paragraph()
    doc.add_heading('推荐实施顺序', level=2)
    doc.add_paragraph('1. 立即可做: 在 Dify 添加 Claude 模型供应商', style='List Number')
    doc.add_paragraph('2. 短期目标: 实现基础 Tool Use (TTM评估)', style='List Number')
    doc.add_paragraph('3. 中期目标: 完善工具生态，集成高级特性', style='List Number')
    doc.add_paragraph('4. 长期目标: 构建完整的智能健康教练系统', style='List Number')

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run('文档版本: v1.0').italic = True
    footer.add_run('\n最后更新: 2026-01-26').italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Save
    output_path = r'D:\behavioral-health-project\docs\CLAUDE_INTEGRATION.docx'
    doc.save(output_path)
    print(f'Word document saved: {output_path}')
    return output_path

if __name__ == '__main__':
    create_word_document()
