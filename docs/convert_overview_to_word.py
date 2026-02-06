"""
Convert PROJECT_OVERVIEW_EXECUTIVE.md to Word format
Executive Summary for Decision Makers
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def create_executive_summary():
    doc = Document()

    # ===== Title =====
    title = doc.add_heading('行为健康平台', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.add_run('项目总览报告').bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run('文档类型: 执行摘要 (Executive Summary)\n').italic = True
    meta.add_run('面向对象: 项目决策者、投资人、管理层\n').italic = True
    meta.add_run('更新日期: 2026-01-26').italic = True
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ===== Section 1: 项目定位 =====
    doc.add_heading('一、项目定位', level=1)

    doc.add_heading('1.1 一句话描述', level=2)
    p = doc.add_paragraph()
    r = p.add_run('行为健康平台是一个基于AI的健康行为改变支持系统，帮助用户养成健康习惯，同时培养专业的健康教练团队。')
    r.bold = True

    doc.add_heading('1.2 解决的核心问题', level=2)

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '痛点'
    hdr[1].text = '现状'
    hdr[2].text = '我们的方案'

    data = [
        ('慢病管理难以持续', '患者缺乏持续指导，行为改变失败率高', 'AI教练7×24小时陪伴支持'),
        ('健康教练供给不足', '专业教练培养周期长、成本高', '标准化培训+AI辅助，快速培养'),
        ('专家知识难以规模化', '专家一对一服务，覆盖人群有限', '专家知识AI化，一对多服务'),
    ]
    for i, (pain, status, solution) in enumerate(data, 1):
        row = table.rows[i].cells
        row[0].text = pain
        row[1].text = status
        row[2].text = solution

    doc.add_paragraph()

    doc.add_heading('1.3 目标用户群', level=2)
    user_table = doc.add_table(rows=2, cols=3)
    user_table.style = 'Table Grid'
    uhdr = user_table.rows[0].cells
    uhdr[0].text = '专家 (知识提供者)'
    uhdr[1].text = '用户/患者 (服务接受者)'
    uhdr[2].text = '健康教练 (服务执行者)'

    udata = user_table.rows[1].cells
    udata[0].text = '• 医生\n• 营养师\n• 运动专家\n• 心理咨询师'
    udata[1].text = '• 慢病患者\n• 亚健康人群\n• 健康管理需求者'
    udata[2].text = '• 在培学员\n• 初级教练\n• 中高级教练\n• 督导专家'

    doc.add_paragraph()

    # ===== Section 2: 产品架构 =====
    doc.add_heading('二、产品架构 ("八爪鱼"模型)', level=1)

    doc.add_heading('2.1 架构说明', level=2)
    doc.add_paragraph('我们采用"八爪鱼"架构，形象地描述系统各部分的关系：')

    arch_table = doc.add_table(rows=6, cols=4)
    arch_table.style = 'Table Grid'
    ahdr = arch_table.rows[0].cells
    ahdr[0].text = '组件'
    ahdr[1].text = '比喻'
    ahdr[2].text = '功能'
    ahdr[3].text = '技术实现'

    arch_data = [
        ('大脑', '章鱼的大脑', 'AI理解和生成能力', 'Ollama本地 + Claude云端'),
        ('数据池', '章鱼的记忆', '健康专业知识库', 'Obsidian + 向量数据库'),
        ('触手1', '专家的手', '专家知识对话服务', 'Dify AI应用平台'),
        ('触手2', '用户的手', '患者行为干预应用', 'H5/小程序'),
        ('触手3', '教练的手', '教练培训管理系统', 'Vue3管理后台'),
    ]
    for i, (comp, metaphor, func, tech) in enumerate(arch_data, 1):
        row = arch_table.rows[i].cells
        row[0].text = comp
        row[1].text = metaphor
        row[2].text = func
        row[3].text = tech

    doc.add_paragraph()

    # ===== Section 3: 三条业务主线 =====
    doc.add_heading('三、三条业务主线详解', level=1)

    doc.add_heading('3.1 触手1: 专家Chatflow', level=2)
    doc.add_paragraph('目标: 让专家的知识通过AI实现规模化服务')
    doc.add_paragraph('当前状态:')
    doc.add_paragraph('• ✅ 已完成: "吃动守恒"专家Chatflow (示范案例)', style='List Bullet')
    doc.add_paragraph('• ⏳ 待开发: 更多专家入驻机制、知识审核流程', style='List Bullet')
    doc.add_paragraph('产品形态: AI对话应用，类似ChatGPT但具备专业健康知识')

    doc.add_heading('3.2 触手2: 用户行为养成', level=2)
    doc.add_paragraph('目标: 帮助用户/患者建立并维持健康行为习惯')
    doc.add_paragraph('核心理论: TTM跨理论模型 (行为改变5阶段)')
    doc.add_paragraph('前意向期 → 意向期 → 准备期 → 行动期 → 维持期')
    doc.add_paragraph('当前状态:')
    doc.add_paragraph('• ✅ 已完成: 理论框架、数据结构设计', style='List Bullet')
    doc.add_paragraph('• ⏳ 待开发: H5用户端应用、行为干预流程', style='List Bullet')
    doc.add_paragraph('产品形态: 移动端H5应用/小程序')

    doc.add_heading('3.3 触手3: 教练培养体系', level=2)
    doc.add_paragraph('目标: 标准化培养健康教练，建立认证体系')
    doc.add_paragraph('教练等级: L0学员 → L1初级 → L2中级 → L3高级 → L4督导')
    doc.add_paragraph('当前状态:')
    doc.add_paragraph('• ✅ 已完成: 登录、课程管理、考试系统、题库、学员管理', style='List Bullet')
    doc.add_paragraph('• 🔨 进行中: 教练管理、晋级审核', style='List Bullet')
    doc.add_paragraph('• ⏳ 待开发: 直播培训、系统设置', style='List Bullet')
    doc.add_paragraph('产品形态: Web管理后台')

    # ===== Section 4: 技术架构 =====
    doc.add_heading('四、技术架构', level=1)

    doc.add_heading('4.1 技术选型', level=2)
    tech_table = doc.add_table(rows=9, cols=3)
    tech_table.style = 'Table Grid'
    thdr = tech_table.rows[0].cells
    thdr[0].text = '层次'
    thdr[1].text = '技术选择'
    thdr[2].text = '选型理由'

    tech_data = [
        ('AI平台', 'Dify', '开源、可视化编排、支持多模型'),
        ('本地模型', 'Ollama + qwen2.5', '免费、隐私保护、中文能力强'),
        ('云端模型', 'Claude API', '复杂推理能力强、工具调用好'),
        ('前端', 'Vue 3 + Ant Design', '成熟生态、组件丰富'),
        ('后端', 'FastAPI', 'Python生态、异步高性能'),
        ('数据库', 'PostgreSQL', '稳定可靠、功能完善'),
        ('向量库', 'Weaviate', '专业向量搜索'),
        ('部署', 'Docker', '环境一致性、便于迁移'),
    ]
    for i, (layer, tech, reason) in enumerate(tech_data, 1):
        row = tech_table.rows[i].cells
        row[0].text = layer
        row[1].text = tech
        row[2].text = reason

    doc.add_paragraph()

    # ===== Section 5: 项目进度 =====
    doc.add_heading('五、项目进度总览', level=1)

    doc.add_heading('5.1 模块完成状态', level=2)

    status_table = doc.add_table(rows=15, cols=4)
    status_table.style = 'Table Grid'
    shdr = status_table.rows[0].cells
    shdr[0].text = '模块'
    shdr[1].text = '所属'
    shdr[2].text = '状态'
    shdr[3].text = '说明'

    status_data = [
        ('Dify平台配置', '基础设施', '✅🔒', '已部署，配置稳定'),
        ('Docker编排', '基础设施', '✅🔒', '8个服务容器正常运行'),
        ('Ollama模型', 'AI层', '✅🔒', 'qwen2.5:0.5b 已配置'),
        ('吃动守恒专家', '触手1', '✅🔒', '首个专家Chatflow完成'),
        ('登录模块', '触手3', '✅🔒', '三种角色权限'),
        ('课程管理', '触手3', '✅', '列表+编辑'),
        ('考试系统', '触手3', '✅', '列表+编辑+成绩'),
        ('题库管理', '触手3', '✅', '列表+编辑'),
        ('学员管理', '触手3', '✅', '含资质上传'),
        ('教练管理', '触手3', '🔨', '开发中'),
        ('晋级审核', '触手3', '🔨', '开发中'),
        ('直播管理', '触手3', '⏳', '待开发'),
        ('用户H5应用', '触手2', '⏳', '待开发'),
        ('FastAPI后端', '业务层', '⏳', '待开发'),
    ]
    for i, (module, belong, status, note) in enumerate(status_data, 1):
        row = status_table.rows[i].cells
        row[0].text = module
        row[1].text = belong
        row[2].text = status
        row[3].text = note

    doc.add_paragraph()

    doc.add_heading('5.2 整体完成度', level=2)
    progress = doc.add_paragraph()
    progress.add_run('基础设施层: 100%\n')
    progress.add_run('AI能力层: 80%\n')
    progress.add_run('触手1(专家): 80%\n')
    progress.add_run('触手2(用户): 20%\n')
    progress.add_run('触手3(教练): 80%\n')
    progress.add_run('后端API层: 20%\n')
    progress.add_run('\n总体进度: 约60%').bold = True

    # ===== Section 6: 下一步工作 =====
    doc.add_heading('六、下一步工作建议', level=1)

    doc.add_heading('6.1 建议开发顺序', level=2)

    phase_table = doc.add_table(rows=6, cols=4)
    phase_table.style = 'Table Grid'
    phdr = phase_table.rows[0].cells
    phdr[0].text = '阶段'
    phdr[1].text = '任务'
    phdr[2].text = '产出'
    phdr[3].text = '里程碑意义'

    phase_data = [
        ('Phase 1', '完成触手3教练管理', '教练列表、详情、晋级审核', '教练培养体系可运行'),
        ('Phase 2', '搭建FastAPI后端', '用户/教练/课程API', '数据可持久化存储'),
        ('Phase 3', '启动触手2用户H5', '用户注册、TTM评估', '用户端可体验'),
        ('Phase 4', '集成Claude API', '智能路由、工具调用', 'AI能力升级'),
        ('Phase 5', '专家入驻机制', '知识上传、审核、收益', '专家生态启动'),
    ]
    for i, (phase, task, output, milestone) in enumerate(phase_data, 1):
        row = phase_table.rows[i].cells
        row[0].text = phase
        row[1].text = task
        row[2].text = output
        row[3].text = milestone

    doc.add_paragraph()

    doc.add_heading('6.2 资源需求评估', level=2)
    resource_table = doc.add_table(rows=6, cols=3)
    resource_table.style = 'Table Grid'
    rhdr = resource_table.rows[0].cells
    rhdr[0].text = '角色'
    rhdr[1].text = '人数'
    rhdr[2].text = '主要职责'

    resource_data = [
        ('前端开发', '1-2人', '管理后台完善、H5开发'),
        ('后端开发', '1-2人', 'FastAPI、数据库设计'),
        ('AI工程师', '1人', '模型调优、RAG优化'),
        ('产品经理', '1人', '需求细化、用户调研'),
        ('测试', '1人', '功能测试、用户体验'),
    ]
    for i, (role, count, duty) in enumerate(resource_data, 1):
        row = resource_table.rows[i].cells
        row[0].text = role
        row[1].text = count
        row[2].text = duty

    doc.add_paragraph()

    # ===== Section 7: 风险与应对 =====
    doc.add_heading('七、风险与应对', level=1)

    risk_table = doc.add_table(rows=5, cols=3)
    risk_table.style = 'Table Grid'
    riskhdr = risk_table.rows[0].cells
    riskhdr[0].text = '风险类型'
    riskhdr[1].text = '具体风险'
    riskhdr[2].text = '应对措施'

    risk_data = [
        ('技术风险', '本地模型能力不足', '已规划Claude API作为补充'),
        ('数据风险', '健康数据敏感', '本地部署、数据加密、权限控制'),
        ('运营风险', '专家资源不足', '先做好示范案例，逐步吸引'),
        ('市场风险', '用户获取困难', '先服务B端，再拓展C端'),
    ]
    for i, (rtype, risk, measure) in enumerate(risk_data, 1):
        row = risk_table.rows[i].cells
        row[0].text = rtype
        row[1].text = risk
        row[2].text = measure

    doc.add_paragraph()

    # ===== Section 8: 项目亮点 =====
    doc.add_heading('八、项目亮点', level=1)

    doc.add_heading('技术亮点', level=2)
    doc.add_paragraph('1. 双引擎AI架构: 本地Ollama保隐私，云端Claude保质量', style='List Number')
    doc.add_paragraph('2. 可视化AI编排: Dify平台让非技术人员也能创建AI应用', style='List Number')
    doc.add_paragraph('3. 模块化设计: 三条触手独立开发，互不影响', style='List Number')
    doc.add_paragraph('4. 变更可控: 冻结机制保护已稳定模块', style='List Number')

    doc.add_heading('业务亮点', level=2)
    doc.add_paragraph('1. 理论支撑: 基于成熟的TTM行为改变模型', style='List Number')
    doc.add_paragraph('2. 闭环设计: 专家→平台→教练→用户→反馈→专家', style='List Number')
    doc.add_paragraph('3. 可扩展性: 支持多专家、多领域健康知识', style='List Number')
    doc.add_paragraph('4. 认证体系: L0-L4五级教练认证，标准化培养', style='List Number')

    # ===== Footer =====
    doc.add_paragraph()
    doc.add_paragraph('─' * 50)
    footer = doc.add_paragraph()
    footer.add_run('GitHub仓库: ').bold = True
    footer.add_run('https://github.com/ICELANF/behavioral-health-project')

    doc.add_paragraph()
    final = doc.add_paragraph()
    final.add_run('本文档持续更新，如有疑问请联系技术规划团队').italic = True
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    output_path = r'D:\behavioral-health-project\docs\PROJECT_OVERVIEW_EXECUTIVE.docx'
    doc.save(output_path)
    print(f'Word document saved: {output_path}')
    return output_path

if __name__ == '__main__':
    create_executive_summary()
