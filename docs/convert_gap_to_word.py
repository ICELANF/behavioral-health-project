"""
Convert GAP_ANALYSIS_BLUEPRINT.md to Word format
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_gap_analysis_doc():
    doc = Document()

    # Title
    title = doc.add_heading('蓝本对照差距分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run('对照文档: 《行为健康数字平台顶层逻辑蓝本（内部统一稿）》\n').italic = True
    meta.add_run('分析日期: 2026-01-26\n').italic = True
    meta.add_run('目的: 识别已实现/未实现功能，提出调整建议').italic = True
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Section 1
    doc.add_heading('一、架构对照：当前 vs 蓝本', level=1)

    doc.add_heading('1.1 架构框架对比', level=2)
    table1 = doc.add_table(rows=5, cols=4)
    table1.style = 'Table Grid'
    hdr = table1.rows[0].cells
    hdr[0].text = '维度'
    hdr[1].text = '当前架构(八爪鱼)'
    hdr[2].text = '蓝本架构'
    hdr[3].text = '差异说明'

    data1 = [
        ('主线划分', '3条触手', '4条主线', '缺少"陪伴追踪主线"'),
        ('层次划分', '5层(L0-L4)', '5层(L1-L5)', '命名不同，内涵有差异'),
        ('核心定位', 'AI健康助手平台', '行为改变操作系统(BCOS)', '蓝本定位更高'),
        ('护城河', '双引擎AI', '双中枢+工作流+知识沉淀', '蓝本更强调系统能力'),
    ]
    for i, row_data in enumerate(data1, 1):
        row = table1.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('关键发现: ').bold = True
    p.add_run('蓝本将"行为评估"独立为第一主线(入口系统)，当前架构未体现。')

    # Section 2
    doc.add_heading('二、四主线×五层系统 实现状态', level=1)

    doc.add_paragraph('图例: ✅已实现  🔨进行中  ⚠️部分  ❌未实现  🔒已冻结')

    matrix_table = doc.add_table(rows=6, cols=5)
    matrix_table.style = 'Table Grid'
    mhdr = matrix_table.rows[0].cells
    mhdr[0].text = '层级'
    mhdr[1].text = 'A行为评估'
    mhdr[2].text = 'B行为干预'
    mhdr[3].text = 'C陪伴追踪'
    mhdr[4].text = 'D成长进化'

    matrix_data = [
        ('L1理念模型层', '⚠️TTM有文档', '⚠️MI有文档', '❌无模型', '⚠️教练等级有'),
        ('L2评估分流层', '❌最大缺口', '❌无规则引擎', '❌无调度系统', '❌无阶段识别'),
        ('L3干预教练层', '❌无问卷系统', '⚠️Prompt在Dify', '❌无策略引擎', '🔨课程/考试'),
        ('L4执行追踪层', '❌无评估H5', '✅🔒专家对话', '❌无用户App', '🔨教练后台'),
        ('L5进化学习层', '❌无模型校准', '❌无效果学习', '❌无策略优化', '❌无路径进化'),
    ]
    for i, row_data in enumerate(matrix_data, 1):
        row = matrix_table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text

    doc.add_paragraph()
    warning = doc.add_paragraph()
    warning.add_run('蓝本明确指出的最大缺口:\n').bold = True
    warning.add_run('❗ A × L2（完整评估分流引擎）\n')
    warning.add_run('❗ C × L2 / L3（陪伴调度与坚持模型）\n')
    warning.add_run('❗ D × L2 / L3（成长阶段与能力模型）')

    # Section 3
    doc.add_heading('三、14大系统实现状态', level=1)

    sys_table = doc.add_table(rows=15, cols=3)
    sys_table.style = 'Table Grid'
    shdr = sys_table.rows[0].cells
    shdr[0].text = '系统'
    shdr[1].text = '状态'
    shdr[2].text = '说明'

    sys_data = [
        ('①行为评估与决策中枢', '❌', '核心缺失-平台第一大脑'),
        ('②知识管理系统', '⚠️', 'Obsidian部分-需结构化'),
        ('③数据基础系统', '❌', '核心缺失'),
        ('④身份与权限中枢', '⚠️', '登录有-权限需完善'),
        ('⑤隐私合规与伦理治理', '❌', '未建设'),
        ('⑥流程编排与工作流', '❌', '核心缺失'),
        ('⑦模型与Agent运维', '⚠️', 'Dify部分'),
        ('⑧监控与质量评估', '❌', '未建设'),
        ('⑨服务输出与陪伴交互', '⚠️', '专家对话有-用户端无'),
        ('⑩专家系统', '⚠️', 'Chatflow有'),
        ('⑪教练与教学系统', '🔨', '后台开发中'),
        ('⑫多业务入口', '⚠️', '仅管理后台'),
        ('⑬动力机制', '❌', '未建设'),
        ('⑭生态与对外接口', '❌', '未建设'),
    ]
    for i, row_data in enumerate(sys_data, 1):
        row = sys_table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text

    doc.add_paragraph()

    # Section 4
    doc.add_heading('四、一期MSS对照', level=1)

    doc.add_paragraph('蓝本要求一期必须优先建设的7个系统:')

    mss_table = doc.add_table(rows=8, cols=4)
    mss_table.style = 'Table Grid'
    msshdr = mss_table.rows[0].cells
    msshdr[0].text = '系统'
    msshdr[1].text = '要求'
    msshdr[2].text = '当前状态'
    msshdr[3].text = '差距'

    mss_data = [
        ('⭐①行为评估与决策中枢', '必须', '❌', '核心缺失'),
        ('⭐②知识管理系统', '必须', '⚠️', 'Obsidian有需结构化'),
        ('⭐③数据基础系统', '必须', '❌', '核心缺失'),
        ('⭐④身份与权限中枢', '必须', '⚠️', '登录有权限需完善'),
        ('⭐⑥流程编排与工作流', '必须', '❌', '核心缺失'),
        ('⭐⑨服务输出与陪伴交互', '必须', '⚠️', '专家对话有用户端无'),
        ('⭐⑪教练与教学系统', '必须', '🔨', '后台开发中'),
    ]
    for i, row_data in enumerate(mss_data, 1):
        row = mss_table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text

    doc.add_paragraph()
    conclusion1 = doc.add_paragraph()
    conclusion1.add_run('结论: ').bold = True
    conclusion1.add_run('7个一期必须系统中，3个核心缺失，3个部分完成，1个开发中。')

    # Section 5
    doc.add_heading('五、开发治理原则对照', level=1)

    rule_table = doc.add_table(rows=5, cols=4)
    rule_table.style = 'Table Grid'
    rulehdr = rule_table.rows[0].cells
    rulehdr[0].text = '原则'
    rulehdr[1].text = '蓝本要求'
    rulehdr[2].text = '当前实践'
    rulehdr[3].text = '评估'

    rule_data = [
        ('原则1', '先挂主线再立模块', '按触手划分', '⚠️需细化'),
        ('原则2', '先建母库再做应用', '直接做页面', '❌违反'),
        ('原则3', '能力组件化', '未组件化', '❌违反'),
        ('原则4', '设立架构官', '无此角色', '❌缺失'),
    ]
    for i, row_data in enumerate(rule_data, 1):
        row = rule_table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text

    doc.add_paragraph()
    order_issue = doc.add_paragraph()
    order_issue.add_run('开发顺序问题:\n').bold = True
    order_issue.add_run('蓝本要求: L1理念 → L2评估 → L3干预 → L4页面 → L5进化\n')
    order_issue.add_run('当前实际: L4页面 → L3部分 → L1部分\n')
    order_issue.add_run('❌ 顺序颠倒').bold = True

    # Section 6
    doc.add_heading('六、核心问题诊断', level=1)

    diag = doc.add_paragraph()
    diag.add_run('蓝本诊断的问题:\n').bold = True
    diag.add_run('❌ 没有"唯一世界观"\n')
    diag.add_run('❌ 没有"唯一主结构"\n')
    diag.add_run('❌ 没有"术语宪法"\n')
    diag.add_run('❌ 没有"架构治理权"\n')
    diag.add_run('\n当前项目状态印证了上述问题。')

    # Section 7
    doc.add_heading('七、调整建议', level=1)

    doc.add_heading('7.1 立即可做 (本周)', level=2)
    doc.add_paragraph('1. 创建统一术语表文档 - 强制使用蓝本术语', style='List Number')
    doc.add_paragraph('2. 创建Trigger Tag字典v1 - 作为评估分流基础', style='List Number')
    doc.add_paragraph('3. 映射现有功能到四主线×五层矩阵', style='List Number')

    doc.add_heading('7.2 下一阶段', level=2)
    doc.add_paragraph('4. 设计行为评估引擎架构 (系统第一大脑)', style='List Number')
    doc.add_paragraph('5. 结构化Obsidian知识库 (系统第二大脑)', style='List Number')
    doc.add_paragraph('6. 建立教练动作母库', style='List Number')

    doc.add_heading('7.3 优先级调整建议', level=2)
    priority = doc.add_paragraph()
    priority.add_run('建议: 暂停L4层新页面开发，优先补齐L2/L3层').bold = True

    # Section 8
    doc.add_heading('八、总结', level=1)

    summary_table = doc.add_table(rows=6, cols=4)
    summary_table.style = 'Table Grid'
    sumhdr = summary_table.rows[0].cells
    sumhdr[0].text = '维度'
    sumhdr[1].text = '蓝本要求'
    sumhdr[2].text = '当前完成'
    sumhdr[3].text = '差距'

    sum_data = [
        ('四主线覆盖', '100%', '50%', '缺A、C主线核心'),
        ('五层系统覆盖', '100%', '30%', 'L2、L5几乎空白'),
        ('14大系统覆盖', '100%', '25%', '7个一期必须缺3个'),
        ('数据体系', '100%', '5%', '几乎未建设'),
        ('术语统一', '100%', '40%', '需建立术语宪法'),
    ]
    for i, row_data in enumerate(sum_data, 1):
        row = summary_table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text

    doc.add_paragraph()

    # Final conclusion
    final = doc.add_paragraph()
    final.add_run('一句话总结:\n').bold = True
    r = final.add_run('当前项目在L4执行层做得较好，但L2评估层(系统大脑)几乎空白，建议暂停页面开发，优先补齐评估分流引擎和知识管理系统，否则各模块难以形成闭环。')
    r.bold = True

    # Save
    output_path = r'D:\behavioral-health-project\docs\GAP_ANALYSIS_BLUEPRINT.docx'
    doc.save(output_path)
    print(f'Word document saved: {output_path}')
    return output_path

if __name__ == '__main__':
    create_gap_analysis_doc()
