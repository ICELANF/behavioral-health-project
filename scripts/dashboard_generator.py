# -*- coding: utf-8 -*-
"""
dashboard_generator.py - 看板生成器

功能:
1. 个人看板 - 展示用户历史评估数据、趋势分析、处方建议
2. 群体看板 - 展示批次统计、风险分布、聚合指标
3. 导出为 Word 文档

使用方法:
    # 生成个人看板
    python scripts/dashboard_generator.py --type individual --user FDBC03D79348

    # 生成群体看板
    python scripts/dashboard_generator.py --type group --batch 2026-01-10

    # 导出为 Word
    python scripts/dashboard_generator.py --type individual --user FDBC03D79348 --export-word
"""

import json
import argparse
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional
from dataclasses import dataclass
from collections import defaultdict
import statistics

# Word 文档生成
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ============ 数据结构 ============

@dataclass
class TrendAnalysis:
    """趋势分析结果"""
    metric_name: str
    current_value: float
    previous_value: float
    change_percent: float
    trend: str  # "improving", "stable", "declining"
    interpretation: str


@dataclass
class IndividualDashboard:
    """个人看板数据"""
    user_id: str
    device_id: str
    generated_at: str

    # 最新评估
    latest_assessment: Dict[str, Any]

    # 历史数据
    assessment_history: List[Dict[str, Any]]
    total_assessments: int

    # 趋势分析
    trends: List[TrendAnalysis]

    # 最新处方
    latest_prescription: Dict[str, Any]

    # 综合评语
    summary: str
    recommendations: List[str]


@dataclass
class GroupDashboard:
    """群体看板数据"""
    batch_name: str
    generated_at: str

    # 基础统计
    total_users: int
    total_assessments: int

    # 风险分布
    risk_distribution: Dict[str, int]  # {"high": 5, "medium": 10, "low": 20}

    # 聚合指标
    avg_composite_score: float
    avg_stress_index: float
    avg_fatigue_index: float
    avg_mood_index: float
    avg_sdnn: float
    avg_rmssd: float

    # 指标分布
    score_distribution: Dict[str, int]  # {"0-40": 2, "41-60": 5, ...}

    # 需重点关注名单
    high_risk_users: List[Dict[str, Any]]

    # 行为模式分布
    behavior_mode_distribution: Dict[str, int]

    # 群体趋势 (如果有多批次数据)
    group_trends: List[Dict[str, Any]]


# ============ 看板生成器 ============

class DashboardGenerator:
    """看板生成器"""

    def __init__(self, project_root: Path = None):
        """
        初始化看板生成器

        Args:
            project_root: 项目根目录
        """
        self.project_root = project_root or Path("D:/behavioral-health-project")
        self.data_dir = self.project_root / "data" / "assessments"
        self.processed_dir = self.data_dir / "processed"
        self.exports_dir = self.data_dir / "exports"

        # 确保目录存在
        (self.exports_dir / "individual").mkdir(parents=True, exist_ok=True)
        (self.exports_dir / "group").mkdir(parents=True, exist_ok=True)

    # ============ 个人看板 ============

    def generate_individual_dashboard(self, user_id: str) -> Optional[IndividualDashboard]:
        """
        生成个人看板

        Args:
            user_id: 用户ID

        Returns:
            IndividualDashboard 对象
        """
        user_dir = self.processed_dir / "users" / user_id
        if not user_dir.exists():
            print(f"[错误] 用户不存在: {user_id}")
            return None

        # 加载所有评估数据
        assessments_dir = user_dir / "assessments"
        assessments = []

        for f in sorted(assessments_dir.glob("*.json")):
            with open(f, 'r', encoding='utf-8') as fp:
                assessments.append(json.load(fp))

        if not assessments:
            print(f"[错误] 用户 {user_id} 无评估数据")
            return None

        # 最新评估
        latest = assessments[-1]

        # 加载最新处方
        rx_dir = user_dir / "prescriptions"
        latest_prescription = {}
        if rx_dir.exists():
            rx_files = sorted(rx_dir.glob("*.json"), reverse=True)
            if rx_files:
                with open(rx_files[0], 'r', encoding='utf-8') as fp:
                    latest_prescription = json.load(fp)

        # 趋势分析
        trends = self._analyze_trends(assessments)

        # 生成综合评语和建议
        summary, recommendations = self._generate_individual_summary(latest, trends)

        # 加载用户档案
        profile_file = user_dir / "profile.json"
        device_id = user_id
        if profile_file.exists():
            with open(profile_file, 'r', encoding='utf-8') as fp:
                profile = json.load(fp)
                device_id = profile.get("device_id", user_id)

        dashboard = IndividualDashboard(
            user_id=user_id,
            device_id=device_id,
            generated_at=datetime.now().isoformat(),
            latest_assessment=latest,
            assessment_history=assessments,
            total_assessments=len(assessments),
            trends=trends,
            latest_prescription=latest_prescription,
            summary=summary,
            recommendations=recommendations
        )

        return dashboard

    def _analyze_trends(self, assessments: List[Dict]) -> List[TrendAnalysis]:
        """分析趋势"""
        trends = []

        if len(assessments) < 2:
            return trends

        current = assessments[-1]
        previous = assessments[-2]

        # 分析各指标趋势
        metrics = [
            ("composite_score", "综合得分", "psych_data"),
            ("stress_index", "压力指数", "psych_data"),
            ("fatigue_index", "疲劳指数", "psych_data"),
            ("mood_index", "心情指数", "psych_data"),
        ]

        for metric_key, metric_name, data_key in metrics:
            current_val = current.get(data_key, {}).get(metric_key, 0)
            previous_val = previous.get(data_key, {}).get(metric_key, 0)

            if previous_val > 0:
                change_pct = ((current_val - previous_val) / previous_val) * 100
            else:
                change_pct = 0

            # 判断趋势
            if metric_key in ["stress_index", "fatigue_index"]:
                # 这些指标降低是好的
                if change_pct < -5:
                    trend = "improving"
                    interpretation = f"{metric_name}下降{abs(change_pct):.1f}%，状态改善"
                elif change_pct > 5:
                    trend = "declining"
                    interpretation = f"{metric_name}上升{change_pct:.1f}%，需要关注"
                else:
                    trend = "stable"
                    interpretation = f"{metric_name}保持稳定"
            else:
                # composite_score, mood_index 升高是好的
                if change_pct > 5:
                    trend = "improving"
                    interpretation = f"{metric_name}提升{change_pct:.1f}%，状态改善"
                elif change_pct < -5:
                    trend = "declining"
                    interpretation = f"{metric_name}下降{abs(change_pct):.1f}%，需要关注"
                else:
                    trend = "stable"
                    interpretation = f"{metric_name}保持稳定"

            trends.append(TrendAnalysis(
                metric_name=metric_name,
                current_value=current_val,
                previous_value=previous_val,
                change_percent=change_pct,
                trend=trend,
                interpretation=interpretation
            ))

        # HRV 趋势
        current_sdnn = current.get("physio_data", {}).get("hrv", {}).get("sdnn", 0)
        previous_sdnn = previous.get("physio_data", {}).get("hrv", {}).get("sdnn", 0)

        if previous_sdnn > 0:
            sdnn_change = ((current_sdnn - previous_sdnn) / previous_sdnn) * 100
            if sdnn_change > 5:
                sdnn_trend = "improving"
                sdnn_interp = f"SDNN提升{sdnn_change:.1f}%，心脏调节能力增强"
            elif sdnn_change < -5:
                sdnn_trend = "declining"
                sdnn_interp = f"SDNN下降{abs(sdnn_change):.1f}%，需注意恢复"
            else:
                sdnn_trend = "stable"
                sdnn_interp = "SDNN保持稳定"

            trends.append(TrendAnalysis(
                metric_name="SDNN (心脏调节)",
                current_value=current_sdnn,
                previous_value=previous_sdnn,
                change_percent=sdnn_change,
                trend=sdnn_trend,
                interpretation=sdnn_interp
            ))

        return trends

    def _generate_individual_summary(
        self,
        latest: Dict,
        trends: List[TrendAnalysis]
    ) -> tuple:
        """生成个人综合评语"""
        psych = latest.get("psych_data", {})
        physio = latest.get("physio_data", {})
        risk = latest.get("risk_assessment", {})

        composite = psych.get("composite_score", 70)
        stress = psych.get("stress_index", 50)
        fatigue = psych.get("fatigue_index", 50)
        mood = psych.get("mood_index", 50)
        sdnn = physio.get("hrv", {}).get("sdnn", 50)

        # 生成综合评语
        if composite >= 80:
            summary = "您当前的整体状态良好，各项指标均处于健康范围内。"
        elif composite >= 60:
            summary = "您当前的状态需要关注，部分指标存在轻度异常。"
        else:
            summary = "您当前的状态需要重点关注，建议及时调整生活方式或寻求专业帮助。"

        # 添加趋势描述
        improving_count = sum(1 for t in trends if t.trend == "improving")
        declining_count = sum(1 for t in trends if t.trend == "declining")

        if improving_count > declining_count:
            summary += " 从趋势来看，您的整体状况在改善中，继续保持。"
        elif declining_count > improving_count:
            summary += " 从趋势来看，您的部分指标有下降趋势，建议引起重视。"

        # 生成建议
        recommendations = []

        if stress > 60:
            recommendations.append("压力指数偏高，建议尝试深呼吸、冥想等放松练习")

        if fatigue > 60:
            recommendations.append("疲劳指数较高，建议保证充足睡眠，适当休息")

        if mood < 50:
            recommendations.append("心情状态欠佳，建议多进行户外活动，与亲友交流")

        if sdnn < 40:
            recommendations.append("心脏调节能力偏低，建议规律运动，避免久坐")

        if not recommendations:
            recommendations.append("保持当前良好的生活习惯")
            recommendations.append("继续定期进行健康评估，跟踪状态变化")

        return summary, recommendations

    # ============ 群体看板 ============

    def generate_group_dashboard(self, batch_name: str = None) -> Optional[GroupDashboard]:
        """
        生成群体看板

        Args:
            batch_name: 批次名称 (可选，默认统计所有数据)

        Returns:
            GroupDashboard 对象
        """
        users_dir = self.processed_dir / "users"
        if not users_dir.exists():
            print("[错误] 无用户数据")
            return None

        # 收集所有评估数据
        all_assessments = []
        user_latest = {}  # 每个用户的最新评估

        for user_dir in users_dir.iterdir():
            if not user_dir.is_dir():
                continue

            assessments_dir = user_dir / "assessments"
            if not assessments_dir.exists():
                continue

            user_id = user_dir.name

            for f in sorted(assessments_dir.glob("*.json")):
                # 如果指定了批次，只统计该批次
                if batch_name:
                    date_str = f.stem  # 2026-01-10
                    if date_str != batch_name:
                        continue

                with open(f, 'r', encoding='utf-8') as fp:
                    data = json.load(fp)
                    data["_user_id"] = user_id
                    all_assessments.append(data)

                    # 记录最新
                    if user_id not in user_latest or f.stem > user_latest[user_id]["assessment_date"]:
                        user_latest[user_id] = data

        if not all_assessments:
            print(f"[错误] 无评估数据 (批次: {batch_name or '全部'})")
            return None

        # 统计分析
        risk_dist = defaultdict(int)
        score_dist = defaultdict(int)
        behavior_dist = defaultdict(int)
        high_risk_users = []

        composite_scores = []
        stress_indices = []
        fatigue_indices = []
        mood_indices = []
        sdnn_values = []
        rmssd_values = []

        for ass in all_assessments:
            psych = ass.get("psych_data", {})
            physio = ass.get("physio_data", {})
            risk = ass.get("risk_assessment", {})

            # 风险分布
            risk_level = risk.get("level", "medium")
            risk_dist[risk_level] += 1

            # 收集高风险用户
            if risk_level == "high":
                high_risk_users.append({
                    "user_id": ass.get("_user_id", ass.get("user_id")),
                    "composite_score": psych.get("composite_score", 0),
                    "risk_flags": risk.get("flags", []),
                    "assessment_date": ass.get("assessment_date", "")
                })

            # 得分分布
            score = psych.get("composite_score", 70)
            if score < 40:
                score_dist["0-40 (高风险)"] += 1
            elif score < 60:
                score_dist["41-60 (需关注)"] += 1
            elif score < 80:
                score_dist["61-80 (一般)"] += 1
            else:
                score_dist["81-100 (良好)"] += 1

            # 收集数值用于计算均值
            composite_scores.append(score)
            stress_indices.append(psych.get("stress_index", 50))
            fatigue_indices.append(psych.get("fatigue_index", 50))
            mood_indices.append(psych.get("mood_index", 50))

            hrv = physio.get("hrv", {})
            if hrv.get("sdnn", 0) > 0:
                sdnn_values.append(hrv["sdnn"])
            if hrv.get("rmssd", 0) > 0:
                rmssd_values.append(hrv["rmssd"])

        # 按综合得分排序高风险用户
        high_risk_users.sort(key=lambda x: x["composite_score"])

        dashboard = GroupDashboard(
            batch_name=batch_name or "全部数据",
            generated_at=datetime.now().isoformat(),
            total_users=len(user_latest),
            total_assessments=len(all_assessments),
            risk_distribution=dict(risk_dist),
            avg_composite_score=statistics.mean(composite_scores) if composite_scores else 0,
            avg_stress_index=statistics.mean(stress_indices) if stress_indices else 0,
            avg_fatigue_index=statistics.mean(fatigue_indices) if fatigue_indices else 0,
            avg_mood_index=statistics.mean(mood_indices) if mood_indices else 0,
            avg_sdnn=statistics.mean(sdnn_values) if sdnn_values else 0,
            avg_rmssd=statistics.mean(rmssd_values) if rmssd_values else 0,
            score_distribution=dict(score_dist),
            high_risk_users=high_risk_users[:10],  # 最多显示10个
            behavior_mode_distribution=dict(behavior_dist),
            group_trends=[]  # 未来可扩展多批次趋势
        )

        return dashboard

    # ============ 导出功能 ============

    def export_individual_to_word(
        self,
        dashboard: IndividualDashboard,
        output_path: str = None
    ) -> Optional[str]:
        """导出个人看板为 Word 文档"""
        if not DOCX_AVAILABLE:
            print("[错误] python-docx 未安装，无法导出 Word")
            return None

        doc = Document()

        # 设置标题
        title = doc.add_heading("个人健康评估报告", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 基本信息
        doc.add_heading("1. 基本信息", level=1)
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'

        info_data = [
            ("用户ID", dashboard.user_id),
            ("设备ID", dashboard.device_id),
            ("评估次数", str(dashboard.total_assessments)),
            ("生成时间", dashboard.generated_at[:10])
        ]

        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = value

        doc.add_paragraph()

        # 最新评估结果
        doc.add_heading("2. 最新评估结果", level=1)

        latest = dashboard.latest_assessment
        psych = latest.get("psych_data", {})
        physio = latest.get("physio_data", {})

        # 心理指标表格
        doc.add_heading("2.1 心理指标", level=2)
        psych_table = doc.add_table(rows=5, cols=3)
        psych_table.style = 'Table Grid'

        psych_headers = ["指标", "数值", "状态"]
        for i, header in enumerate(psych_headers):
            psych_table.rows[0].cells[i].text = header

        psych_data = [
            ("综合得分", psych.get("composite_score", 0), self._score_status(psych.get("composite_score", 0))),
            ("压力指数", psych.get("stress_index", 0), self._index_status(psych.get("stress_index", 0), reverse=True)),
            ("疲劳指数", psych.get("fatigue_index", 0), self._index_status(psych.get("fatigue_index", 0), reverse=True)),
            ("心情指数", psych.get("mood_index", 0), self._index_status(psych.get("mood_index", 0)))
        ]

        for i, (metric, value, status) in enumerate(psych_data):
            psych_table.rows[i+1].cells[0].text = metric
            psych_table.rows[i+1].cells[1].text = f"{value:.0f}"
            psych_table.rows[i+1].cells[2].text = status

        doc.add_paragraph()

        # 生理指标表格
        doc.add_heading("2.2 生理指标", level=2)
        physio_table = doc.add_table(rows=4, cols=3)
        physio_table.style = 'Table Grid'

        for i, header in enumerate(["指标", "数值", "说明"]):
            physio_table.rows[0].cells[i].text = header

        hrv = physio.get("hrv", {})
        hr = physio.get("heart_rate", {})

        physio_data = [
            ("平均心率", f"{hr.get('avg', 0):.0f} bpm", "正常范围 60-100"),
            ("SDNN (心脏调节)", f"{hrv.get('sdnn', 0):.0f} ms", "建议 > 50ms"),
            ("RMSSD (恢复能力)", f"{hrv.get('rmssd', 0):.0f} ms", "建议 > 30ms")
        ]

        for i, (metric, value, desc) in enumerate(physio_data):
            physio_table.rows[i+1].cells[0].text = metric
            physio_table.rows[i+1].cells[1].text = value
            physio_table.rows[i+1].cells[2].text = desc

        doc.add_paragraph()

        # 趋势分析
        if dashboard.trends:
            doc.add_heading("3. 趋势分析", level=1)

            for trend in dashboard.trends:
                trend_icon = "↑" if trend.trend == "improving" else ("↓" if trend.trend == "declining" else "→")
                p = doc.add_paragraph()
                p.add_run(f"{trend_icon} {trend.metric_name}: ").bold = True
                p.add_run(trend.interpretation)

        doc.add_paragraph()

        # 综合评语
        doc.add_heading("4. 综合评语", level=1)
        doc.add_paragraph(dashboard.summary)

        # 建议
        doc.add_heading("5. 健康建议", level=1)
        for rec in dashboard.recommendations:
            doc.add_paragraph(f"• {rec}")

        # 处方摘要
        if dashboard.latest_prescription:
            doc.add_heading("6. 行为处方", level=1)
            rx = dashboard.latest_prescription

            rx_meta = rx.get("prescription_meta", {})
            doc.add_paragraph(f"处方名称: {rx_meta.get('name', '未知')}")
            doc.add_paragraph(f"干预策略: {rx_meta.get('intervention_strategy', '未知')}")

            if rx.get("tasks"):
                doc.add_heading("6.1 推荐任务", level=2)
                for task in rx.get("tasks", [])[:3]:
                    doc.add_paragraph(f"• {task.get('name', '')}: {task.get('description', '')}")

        # 页脚
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("—— 行健行为教练 · 行为健康促进系统 ——").italic = True

        # 保存文件
        if not output_path:
            output_path = self.exports_dir / "individual" / f"{dashboard.user_id}_{datetime.now().strftime('%Y-%m-%d')}.docx"

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        print(f"[导出] 个人看板已保存: {output_path}")
        return str(output_path)

    def export_group_to_word(
        self,
        dashboard: GroupDashboard,
        output_path: str = None
    ) -> Optional[str]:
        """导出群体看板为 Word 文档"""
        if not DOCX_AVAILABLE:
            print("[错误] python-docx 未安装，无法导出 Word")
            return None

        doc = Document()

        # 设置标题
        title = doc.add_heading("群体健康评估报告", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run(f"批次: {dashboard.batch_name}").bold = True

        # 概览
        doc.add_heading("1. 数据概览", level=1)
        overview_table = doc.add_table(rows=3, cols=2)
        overview_table.style = 'Table Grid'

        overview_data = [
            ("参与人数", str(dashboard.total_users)),
            ("评估总数", str(dashboard.total_assessments)),
            ("生成时间", dashboard.generated_at[:10])
        ]

        for i, (label, value) in enumerate(overview_data):
            overview_table.rows[i].cells[0].text = label
            overview_table.rows[i].cells[1].text = value

        doc.add_paragraph()

        # 风险分布
        doc.add_heading("2. 风险分布", level=1)
        risk_table = doc.add_table(rows=4, cols=3)
        risk_table.style = 'Table Grid'

        risk_table.rows[0].cells[0].text = "风险等级"
        risk_table.rows[0].cells[1].text = "人数"
        risk_table.rows[0].cells[2].text = "占比"

        total = sum(dashboard.risk_distribution.values()) or 1
        risk_labels = {"high": "高风险", "medium": "中等风险", "low": "低风险"}

        for i, (level, label) in enumerate(risk_labels.items()):
            count = dashboard.risk_distribution.get(level, 0)
            risk_table.rows[i+1].cells[0].text = label
            risk_table.rows[i+1].cells[1].text = str(count)
            risk_table.rows[i+1].cells[2].text = f"{count/total*100:.1f}%"

        doc.add_paragraph()

        # 得分分布
        doc.add_heading("3. 得分分布", level=1)
        score_table = doc.add_table(rows=len(dashboard.score_distribution)+1, cols=3)
        score_table.style = 'Table Grid'

        score_table.rows[0].cells[0].text = "得分区间"
        score_table.rows[0].cells[1].text = "人数"
        score_table.rows[0].cells[2].text = "占比"

        for i, (range_name, count) in enumerate(sorted(dashboard.score_distribution.items())):
            score_table.rows[i+1].cells[0].text = range_name
            score_table.rows[i+1].cells[1].text = str(count)
            score_table.rows[i+1].cells[2].text = f"{count/total*100:.1f}%"

        doc.add_paragraph()

        # 群体平均指标
        doc.add_heading("4. 群体平均指标", level=1)
        avg_table = doc.add_table(rows=7, cols=2)
        avg_table.style = 'Table Grid'

        avg_data = [
            ("平均综合得分", f"{dashboard.avg_composite_score:.1f}"),
            ("平均压力指数", f"{dashboard.avg_stress_index:.1f}"),
            ("平均疲劳指数", f"{dashboard.avg_fatigue_index:.1f}"),
            ("平均心情指数", f"{dashboard.avg_mood_index:.1f}"),
            ("平均 SDNN", f"{dashboard.avg_sdnn:.1f} ms"),
            ("平均 RMSSD", f"{dashboard.avg_rmssd:.1f} ms")
        ]

        avg_table.rows[0].cells[0].text = "指标"
        avg_table.rows[0].cells[1].text = "群体均值"

        for i, (metric, value) in enumerate(avg_data):
            avg_table.rows[i+1].cells[0].text = metric
            avg_table.rows[i+1].cells[1].text = value

        doc.add_paragraph()

        # 高风险人员名单
        if dashboard.high_risk_users:
            doc.add_heading("5. 重点关注名单", level=1)
            doc.add_paragraph("以下人员需重点关注和干预:")

            hr_table = doc.add_table(rows=len(dashboard.high_risk_users)+1, cols=4)
            hr_table.style = 'Table Grid'

            hr_table.rows[0].cells[0].text = "用户ID"
            hr_table.rows[0].cells[1].text = "综合得分"
            hr_table.rows[0].cells[2].text = "评估日期"
            hr_table.rows[0].cells[3].text = "风险提示"

            for i, user in enumerate(dashboard.high_risk_users):
                hr_table.rows[i+1].cells[0].text = user["user_id"][-4:]  # 只显示后4位
                hr_table.rows[i+1].cells[1].text = f"{user['composite_score']:.0f}"
                hr_table.rows[i+1].cells[2].text = user["assessment_date"]
                hr_table.rows[i+1].cells[3].text = "; ".join(user["risk_flags"][:2]) if user["risk_flags"] else "-"

        # 建议
        doc.add_heading("6. 群体干预建议", level=1)

        suggestions = []
        if dashboard.avg_stress_index > 60:
            suggestions.append("群体压力偏高，建议开展减压团体活动")
        if dashboard.avg_fatigue_index > 60:
            suggestions.append("群体疲劳明显，建议调整工作节奏，增加休息时间")
        if dashboard.avg_composite_score < 70:
            suggestions.append("群体整体健康状况需改善，建议加强健康宣教")
        if len(dashboard.high_risk_users) > dashboard.total_users * 0.2:
            suggestions.append("高风险人群占比较高，建议对重点人员进行一对一跟进")

        if not suggestions:
            suggestions.append("群体整体状况良好，继续保持当前健康管理措施")

        for sug in suggestions:
            doc.add_paragraph(f"• {sug}")

        # 页脚
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("—— 行健行为教练 · 行为健康促进系统 ——").italic = True

        # 保存文件
        if not output_path:
            batch_safe = dashboard.batch_name.replace("/", "-").replace("\\", "-")
            output_path = self.exports_dir / "group" / f"batch_{batch_safe}_summary.docx"

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        print(f"[导出] 群体看板已保存: {output_path}")
        return str(output_path)

    def _score_status(self, score: float) -> str:
        """根据得分返回状态"""
        if score >= 80:
            return "良好"
        elif score >= 60:
            return "一般"
        else:
            return "需关注"

    def _index_status(self, index: float, reverse: bool = False) -> str:
        """根据指数返回状态"""
        if reverse:
            # 压力、疲劳等，低为好
            if index < 40:
                return "良好"
            elif index < 60:
                return "一般"
            else:
                return "偏高"
        else:
            # 心情等，高为好
            if index >= 60:
                return "良好"
            elif index >= 40:
                return "一般"
            else:
                return "偏低"

    # ============ JSON 导出 ============

    def export_individual_to_json(self, dashboard: IndividualDashboard) -> str:
        """导出个人看板为 JSON"""
        data = {
            "user_id": dashboard.user_id,
            "device_id": dashboard.device_id,
            "generated_at": dashboard.generated_at,
            "total_assessments": dashboard.total_assessments,
            "latest_assessment": dashboard.latest_assessment,
            "trends": [
                {
                    "metric_name": t.metric_name,
                    "current_value": t.current_value,
                    "previous_value": t.previous_value,
                    "change_percent": t.change_percent,
                    "trend": t.trend,
                    "interpretation": t.interpretation
                }
                for t in dashboard.trends
            ],
            "summary": dashboard.summary,
            "recommendations": dashboard.recommendations,
            "latest_prescription": dashboard.latest_prescription
        }

        output_path = self.exports_dir / "individual" / f"{dashboard.user_id}_dashboard.json"
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

        print(f"[导出] JSON 已保存: {output_path}")
        return str(output_path)

    def export_group_to_json(self, dashboard: GroupDashboard) -> str:
        """导出群体看板为 JSON"""
        data = {
            "batch_name": dashboard.batch_name,
            "generated_at": dashboard.generated_at,
            "total_users": dashboard.total_users,
            "total_assessments": dashboard.total_assessments,
            "risk_distribution": dashboard.risk_distribution,
            "score_distribution": dashboard.score_distribution,
            "avg_composite_score": dashboard.avg_composite_score,
            "avg_stress_index": dashboard.avg_stress_index,
            "avg_fatigue_index": dashboard.avg_fatigue_index,
            "avg_mood_index": dashboard.avg_mood_index,
            "avg_sdnn": dashboard.avg_sdnn,
            "avg_rmssd": dashboard.avg_rmssd,
            "high_risk_users": dashboard.high_risk_users,
            "behavior_mode_distribution": dashboard.behavior_mode_distribution
        }

        batch_safe = dashboard.batch_name.replace("/", "-").replace("\\", "-")
        output_path = self.exports_dir / "group" / f"batch_{batch_safe}.json"
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

        print(f"[导出] JSON 已保存: {output_path}")
        return str(output_path)


# ============ 打印功能 ============

def print_individual_dashboard(dashboard: IndividualDashboard):
    """打印个人看板到控制台"""
    print("\n" + "=" * 60)
    print("个人健康评估看板")
    print("=" * 60)

    print(f"\n用户ID: {dashboard.user_id}")
    print(f"设备ID: {dashboard.device_id}")
    print(f"评估次数: {dashboard.total_assessments}")

    # 最新评估
    latest = dashboard.latest_assessment
    psych = latest.get("psych_data", {})
    physio = latest.get("physio_data", {})

    print(f"\n--- 最新评估 ({latest.get('assessment_date', 'N/A')}) ---")
    print(f"  综合得分: {psych.get('composite_score', 0):.0f}")
    print(f"  压力指数: {psych.get('stress_index', 0):.0f}")
    print(f"  疲劳指数: {psych.get('fatigue_index', 0):.0f}")
    print(f"  心情指数: {psych.get('mood_index', 0):.0f}")
    print(f"  SDNN: {physio.get('hrv', {}).get('sdnn', 0):.0f} ms")
    print(f"  RMSSD: {physio.get('hrv', {}).get('rmssd', 0):.0f} ms")

    # 趋势
    if dashboard.trends:
        print(f"\n--- 趋势分析 ---")
        for trend in dashboard.trends:
            icon = "↑" if trend.trend == "improving" else ("↓" if trend.trend == "declining" else "→")
            print(f"  {icon} {trend.interpretation}")

    # 评语
    print(f"\n--- 综合评语 ---")
    print(f"  {dashboard.summary}")

    # 建议
    print(f"\n--- 健康建议 ---")
    for rec in dashboard.recommendations:
        print(f"  • {rec}")

    print("\n" + "=" * 60)


def print_group_dashboard(dashboard: GroupDashboard):
    """打印群体看板到控制台"""
    print("\n" + "=" * 60)
    print("群体健康评估看板")
    print("=" * 60)

    print(f"\n批次: {dashboard.batch_name}")
    print(f"参与人数: {dashboard.total_users}")
    print(f"评估总数: {dashboard.total_assessments}")

    # 风险分布
    print(f"\n--- 风险分布 ---")
    total = sum(dashboard.risk_distribution.values()) or 1
    for level, count in dashboard.risk_distribution.items():
        label = {"high": "高风险", "medium": "中等风险", "low": "低风险"}.get(level, level)
        print(f"  {label}: {count} ({count/total*100:.1f}%)")

    # 得分分布
    print(f"\n--- 得分分布 ---")
    for range_name, count in sorted(dashboard.score_distribution.items()):
        print(f"  {range_name}: {count}")

    # 平均指标
    print(f"\n--- 群体均值 ---")
    print(f"  平均综合得分: {dashboard.avg_composite_score:.1f}")
    print(f"  平均压力指数: {dashboard.avg_stress_index:.1f}")
    print(f"  平均疲劳指数: {dashboard.avg_fatigue_index:.1f}")
    print(f"  平均心情指数: {dashboard.avg_mood_index:.1f}")
    print(f"  平均 SDNN: {dashboard.avg_sdnn:.1f} ms")
    print(f"  平均 RMSSD: {dashboard.avg_rmssd:.1f} ms")

    # 高风险名单
    if dashboard.high_risk_users:
        print(f"\n--- 重点关注名单 ---")
        for user in dashboard.high_risk_users[:5]:
            print(f"  • {user['user_id'][-4:]}: 得分={user['composite_score']:.0f}, 日期={user['assessment_date']}")

    print("\n" + "=" * 60)


# ============ 主函数 ============

def main():
    parser = argparse.ArgumentParser(description="看板生成器")
    parser.add_argument("--type", choices=["individual", "group"], required=True, help="看板类型")
    parser.add_argument("--user", help="用户ID (个人看板必需)")
    parser.add_argument("--batch", help="批次名称 (群体看板可选)")
    parser.add_argument("--export-word", action="store_true", help="导出为 Word 文档")
    parser.add_argument("--export-json", action="store_true", help="导出为 JSON")
    parser.add_argument("--quiet", action="store_true", help="不打印到控制台")

    args = parser.parse_args()

    generator = DashboardGenerator()

    if args.type == "individual":
        if not args.user:
            print("[错误] 个人看板需要指定 --user")
            return

        dashboard = generator.generate_individual_dashboard(args.user)
        if not dashboard:
            return

        if not args.quiet:
            print_individual_dashboard(dashboard)

        if args.export_word:
            generator.export_individual_to_word(dashboard)

        if args.export_json:
            generator.export_individual_to_json(dashboard)

    elif args.type == "group":
        dashboard = generator.generate_group_dashboard(args.batch)
        if not dashboard:
            return

        if not args.quiet:
            print_group_dashboard(dashboard)

        if args.export_word:
            generator.export_group_to_word(dashboard)

        if args.export_json:
            generator.export_group_to_json(dashboard)


if __name__ == "__main__":
    main()
