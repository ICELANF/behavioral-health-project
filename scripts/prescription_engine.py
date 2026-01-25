# -*- coding: utf-8 -*-
"""
prescription_engine.py - 行为处方引擎

功能:
1. BehavioralProfile 类: 行为倾向、模式、动机强度、改变阶段计算
2. generate_prescription 函数: 基于生理/心理指标生成处方
3. 集成输出: Agent 可读的 JSON 结构

规则逻辑:
- SDNN < 30 且 疲劳度 > 70 → 机体耗竭模式 → 强制性修复处方
- 精力 < 40 → 微习惯启动策略 (而非高强度目标)
"""

import json
import uuid
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, Any, Optional, List, Tuple
from dataclasses import dataclass, field, asdict
from enum import Enum

# Word 文档导出支持
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ============ 枚举定义 ============

class BehaviorMode(Enum):
    """行为模式枚举"""
    EXHAUSTION = "exhaustion"           # 机体耗竭模式
    HIDDEN_FATIGUE = "hidden_fatigue"   # 隐性疲劳模式 (HRV恢复能力低但主观压力低)
    OVERCOMPENSATION = "overcompensation"  # 过度补偿模式
    STRESS_AVOIDANCE = "stress_avoidance"  # 应激逃避模式
    SOMATIZATION = "somatization"       # 躯体化模式
    EMOTIONAL_DYSREGULATION = "emotional_dysregulation"  # 情绪失调模式
    BALANCED = "balanced"               # 平衡模式


class MotivationLevel(Enum):
    """动机强度等级"""
    DEPLETED = "depleted"       # 耗竭 (< 30)
    LOW = "low"                 # 低 (30-50)
    MODERATE = "moderate"       # 中等 (50-70)
    HIGH = "high"               # 高 (> 70)


class ChangeStage(Enum):
    """改变阶段"""
    PRECONTEMPLATION = "precontemplation"  # 前意向期
    CONTEMPLATION = "contemplation"        # 意向期
    PREPARATION = "preparation"            # 准备期
    ACTION = "action"                      # 行动期
    MAINTENANCE = "maintenance"            # 维持期


class InterventionStrategy(Enum):
    """干预策略"""
    MANDATORY_RECOVERY = "mandatory_recovery"   # 强制性修复
    HIDDEN_FATIGUE_RECOVERY = "hidden_fatigue_recovery"  # 隐性疲劳修复
    MICRO_HABIT = "micro_habit"                 # 微习惯启动
    GRADUAL_BUILDUP = "gradual_buildup"         # 渐进式提升
    HABIT_STRENGTHENING = "habit_strengthening" # 习惯强化
    IDENTITY_CONSOLIDATION = "identity_consolidation"  # 身份巩固


# ============ BehavioralProfile 类 ============

@dataclass
class BehavioralProfile:
    """
    行为画像类

    包含行为倾向、模式、动机强度和改变阶段的计算逻辑
    """
    # 原始指标
    sdnn: float = 0.0
    rmssd: float = 0.0
    heart_rate: float = 75.0
    fatigue_index: float = 50.0
    energy_level: float = 50.0
    mood_score: float = 50.0
    anxiety_score: float = 30.0
    stress_index: float = 50.0
    sleep_quality: float = 70.0

    # 历史数据
    assessment_count: int = 1
    consecutive_improvement_weeks: int = 0
    task_completion_rate: float = 0.0

    # 计算结果
    behavior_mode: BehaviorMode = BehaviorMode.BALANCED
    motivation_level: MotivationLevel = MotivationLevel.MODERATE
    change_stage: ChangeStage = ChangeStage.CONTEMPLATION
    intervention_strategy: InterventionStrategy = InterventionStrategy.GRADUAL_BUILDUP

    # 诊断标志
    is_exhaustion_mode: bool = False
    is_hidden_fatigue_mode: bool = False  # 隐性疲劳模式标志
    is_low_motivation: bool = False
    risk_level: str = "medium"

    # 计算分数
    motivation_score: float = 50.0
    energy_mood_match: float = 0.5
    spi_coefficient: float = 0.7

    def __post_init__(self):
        """初始化后自动计算各项指标"""
        self.calculate_all()

    def calculate_all(self):
        """计算所有指标"""
        self._calculate_behavior_mode()
        self._calculate_motivation()
        self._calculate_change_stage()
        self._determine_intervention_strategy()
        self._assess_risk_level()

    def _calculate_behavior_mode(self):
        """
        计算行为模式

        规则1: SDNN < 30 且 疲劳度 > 70 → 机体耗竭模式
        规则1.5: HRV恢复能力低 (SDNN < 50) 且 主观压力低 (stress_index < 50) → 隐性疲劳模式
        """
        # 规则1: 机体耗竭模式检测
        if self.sdnn < 30 and self.fatigue_index > 70:
            self.behavior_mode = BehaviorMode.EXHAUSTION
            self.is_exhaustion_mode = True
            return

        # 规则1.5: 隐性疲劳模式检测
        # HRV恢复能力低但主观压力低 → 身体已疲劳但心理尚未察觉
        hrv_recovery_low = self.sdnn < 50 or self.rmssd < 30
        subjective_stress_low = self.stress_index < 50 and self.anxiety_score < 50
        if hrv_recovery_low and subjective_stress_low and self.fatigue_index < 60:
            self.behavior_mode = BehaviorMode.HIDDEN_FATIGUE
            self.is_hidden_fatigue_mode = True
            return

        # 规则2: 过度补偿模式
        if self.sdnn < 40 and self.anxiety_score > 60 and self.energy_level > 60:
            self.behavior_mode = BehaviorMode.OVERCOMPENSATION
            return

        # 规则3: 应激逃避模式
        if self.stress_index > 70 and self.energy_level < 40:
            self.behavior_mode = BehaviorMode.STRESS_AVOIDANCE
            return

        # 规则4: 躯体化模式
        if self.sdnn < 30 and self.fatigue_index > 60 and self.anxiety_score < 50:
            self.behavior_mode = BehaviorMode.SOMATIZATION
            return

        # 规则5: 情绪失调模式
        mood_variability = abs(self.energy_level - self.mood_score)
        if mood_variability > 30 and self.anxiety_score > 50:
            self.behavior_mode = BehaviorMode.EMOTIONAL_DYSREGULATION
            return

        # 默认: 平衡模式
        self.behavior_mode = BehaviorMode.BALANCED

    def _calculate_motivation(self):
        """
        计算动机强度

        规则: 精力 < 40 → 低动机状态，建议微习惯启动策略

        计算公式:
        1. energy_mood_match = 1 - abs(energy - mood) / 100
        2. base_motivation = (energy + mood) / 2
        3. motivation = base_motivation * (0.5 + 0.5 * match_score)
        """
        # 计算匹配度
        self.energy_mood_match = 1 - abs(self.energy_level - self.mood_score) / 100

        # 计算基础动机
        base_motivation = (self.energy_level + self.mood_score) / 2

        # 计算最终动机分数
        self.motivation_score = base_motivation * (0.5 + 0.5 * self.energy_mood_match)

        # 判定动机等级
        if self.energy_level < 40 or self.motivation_score < 30:
            self.motivation_level = MotivationLevel.DEPLETED
            self.is_low_motivation = True
        elif self.motivation_score < 50:
            self.motivation_level = MotivationLevel.LOW
            self.is_low_motivation = True
        elif self.motivation_score < 70:
            self.motivation_level = MotivationLevel.MODERATE
        else:
            self.motivation_level = MotivationLevel.HIGH

    def _calculate_change_stage(self):
        """
        计算改变阶段

        基于历史数据频率判定:
        - 初次测评 = 意向期
        - 连续好转 = 行动期
        """
        if self.assessment_count <= 1:
            self.change_stage = ChangeStage.CONTEMPLATION
            self.spi_coefficient = 0.5
        elif self.consecutive_improvement_weeks >= 8 and self.task_completion_rate >= 0.8:
            self.change_stage = ChangeStage.MAINTENANCE
            self.spi_coefficient = 1.0
        elif self.consecutive_improvement_weeks >= 2:
            self.change_stage = ChangeStage.ACTION
            self.spi_coefficient = 0.9
        elif self.task_completion_rate > 0.3:
            self.change_stage = ChangeStage.PREPARATION
            self.spi_coefficient = 0.7
        else:
            self.change_stage = ChangeStage.CONTEMPLATION
            self.spi_coefficient = 0.5

    def _determine_intervention_strategy(self):
        """
        确定干预策略

        规则:
        1. 机体耗竭模式 → 强制性修复处方
        1.5. 隐性疲劳模式 → 隐性疲劳修复处方
        2. 精力 < 40 → 微习惯启动策略
        3. 其他根据改变阶段匹配
        """
        # 规则1: 机体耗竭 → 强制性修复
        if self.is_exhaustion_mode:
            self.intervention_strategy = InterventionStrategy.MANDATORY_RECOVERY
            return

        # 规则1.5: 隐性疲劳 → 隐性疲劳修复
        if self.is_hidden_fatigue_mode:
            self.intervention_strategy = InterventionStrategy.HIDDEN_FATIGUE_RECOVERY
            return

        # 规则2: 低动机 → 微习惯启动
        if self.is_low_motivation or self.energy_level < 40:
            self.intervention_strategy = InterventionStrategy.MICRO_HABIT
            return

        # 规则3: 根据改变阶段匹配
        stage_strategy_map = {
            ChangeStage.PRECONTEMPLATION: InterventionStrategy.MICRO_HABIT,
            ChangeStage.CONTEMPLATION: InterventionStrategy.MICRO_HABIT,
            ChangeStage.PREPARATION: InterventionStrategy.GRADUAL_BUILDUP,
            ChangeStage.ACTION: InterventionStrategy.HABIT_STRENGTHENING,
            ChangeStage.MAINTENANCE: InterventionStrategy.IDENTITY_CONSOLIDATION
        }
        self.intervention_strategy = stage_strategy_map.get(
            self.change_stage,
            InterventionStrategy.GRADUAL_BUILDUP
        )

    def _assess_risk_level(self):
        """评估风险等级"""
        risk_score = 0

        if self.sdnn < 30:
            risk_score += 3
        elif self.sdnn < 50:
            risk_score += 1

        if self.fatigue_index > 70:
            risk_score += 2

        if self.anxiety_score > 60:
            risk_score += 2

        if self.energy_level < 30:
            risk_score += 2

        if self.sleep_quality < 50:
            risk_score += 1

        if risk_score >= 6:
            self.risk_level = "high"
        elif risk_score >= 3:
            self.risk_level = "medium"
        else:
            self.risk_level = "low"

    def to_dict(self) -> Dict[str, Any]:
        """转换为字典"""
        return {
            "raw_metrics": {
                "sdnn": self.sdnn,
                "rmssd": self.rmssd,
                "heart_rate": self.heart_rate,
                "fatigue_index": self.fatigue_index,
                "energy_level": self.energy_level,
                "mood_score": self.mood_score,
                "anxiety_score": self.anxiety_score,
                "stress_index": self.stress_index,
                "sleep_quality": self.sleep_quality
            },
            "history": {
                "assessment_count": self.assessment_count,
                "consecutive_improvement_weeks": self.consecutive_improvement_weeks,
                "task_completion_rate": self.task_completion_rate
            },
            "analysis": {
                "behavior_mode": self.behavior_mode.value,
                "motivation_level": self.motivation_level.value,
                "change_stage": self.change_stage.value,
                "intervention_strategy": self.intervention_strategy.value,
                "risk_level": self.risk_level
            },
            "scores": {
                "motivation_score": round(self.motivation_score, 1),
                "energy_mood_match": round(self.energy_mood_match, 2),
                "spi_coefficient": self.spi_coefficient
            },
            "flags": {
                "is_exhaustion_mode": self.is_exhaustion_mode,
                "is_hidden_fatigue_mode": self.is_hidden_fatigue_mode,
                "is_low_motivation": self.is_low_motivation
            }
        }

    @classmethod
    def from_extraction_result(cls, physio_data: Dict, psych_data: Dict,
                                history: Dict = None) -> 'BehavioralProfile':
        """从数据提取结果创建 BehavioralProfile"""
        history = history or {}

        return cls(
            sdnn=physio_data.get('SDNN', 50.0),
            rmssd=physio_data.get('RMSSD', 30.0),
            heart_rate=physio_data.get('heart_rate', 75.0),
            fatigue_index=psych_data.get('fatigue_index', 50.0),
            energy_level=psych_data.get('energy', psych_data.get('energy_level', 50.0)),
            mood_score=psych_data.get('mood', psych_data.get('mood_score', 50.0)),
            anxiety_score=psych_data.get('anxiety_score', 30.0),
            stress_index=psych_data.get('stress_index', 50.0),
            sleep_quality=physio_data.get('sleep_quality', 70.0),
            assessment_count=history.get('assessment_count', 1),
            consecutive_improvement_weeks=history.get('consecutive_improvement_weeks', 0),
            task_completion_rate=history.get('task_completion_rate', 0.0)
        )


# ============ 处方模板 ============

PRESCRIPTION_TEMPLATES = {
    InterventionStrategy.MANDATORY_RECOVERY: {
        "name": "强制性修复处方",
        "description": "针对机体耗竭模式，优先恢复生理基础",
        "priority": "urgent",
        "max_tasks": 2,
        "max_difficulty": 1,
        "tone": "gentle_protective",
        "script": {
            "opening": "您的身体正在发出强烈的休息信号。现在最重要的不是做更多，而是让身体得到真正的恢复。",
            "motivation": "身体的恢复是一切改变的基础。没有能量，任何目标都难以实现。先照顾好自己。",
            "closing": "这周只有一个任务：好好休息。其他的，等身体恢复了再说。"
        },
        "tasks": [
            {
                "content": "确保每晚7-8小时睡眠",
                "description": "设定固定的睡眠时间，创造黑暗安静的睡眠环境",
                "difficulty": 1,
                "type": "recovery",
                "duration_minutes": 0,
                "frequency": "每日"
            },
            {
                "content": "每天进行10分钟深呼吸放松",
                "description": "找一个安静的地方，进行腹式呼吸，让身心放松",
                "difficulty": 1,
                "type": "recovery",
                "duration_minutes": 10,
                "frequency": "每日2次"
            }
        ],
        "knowledge": [
            {
                "knowledge_id": "KP-RECOVERY-001",
                "title": "身体耗竭的信号与恢复",
                "summary": "了解身体疲劳的生理机制，学会识别耗竭信号，掌握科学恢复方法。"
            }
        ],
        "videos": [
            {
                "video_id": "VID-RECOVERY-001",
                "title": "深度放松引导冥想",
                "description": "15分钟身体放松引导，帮助快速进入休息状态",
                "duration_seconds": 900
            }
        ],
        "products": [
            {
                "product_id": "PROD-RECOVERY-001",
                "name": "助眠香薰套装",
                "relevance": "帮助创造放松的睡眠环境"
            }
        ],
        "warnings": [
            "暂停所有高强度运动",
            "减少咖啡因摄入",
            "避免熬夜和过度工作",
            "如症状持续，建议就医检查"
        ]
    },

    InterventionStrategy.HIDDEN_FATIGUE_RECOVERY: {
        "name": "隐性疲劳修复处方",
        "description": "针对HRV恢复能力低但主观压力感知低的隐性疲劳状态",
        "priority": "high",
        "max_tasks": 2,
        "max_difficulty": 1,
        "tone": "gentle_informative",
        "script": {
            "opening": "您的身体数据显示存在隐性疲劳——虽然您可能感觉还好，但HRV指标提示身体的恢复储备已经偏低。这是身体在发出早期预警信号。",
            "motivation": "识别隐性疲劳是非常重要的自我觉察能力。现在调整，可以避免身体进入真正的耗竭状态。预防胜于治疗。",
            "closing": "这周的重点是提升身体的恢复能力。通过科学的休息和监测，让您的HRV恢复储备回到健康水平。"
        },
        "tasks": [
            {
                "content": "每日进行HRV恢复监测",
                "description": "使用穿戴设备记录晨起HRV，追踪恢复趋势",
                "difficulty": 1,
                "type": "monitoring",
                "duration_minutes": 5,
                "frequency": "每日晨起"
            },
            {
                "content": "午间主动休息15分钟",
                "description": "中午安排一次短暂休息，可以是小睡或放松呼吸",
                "difficulty": 1,
                "type": "recovery",
                "duration_minutes": 15,
                "frequency": "每日"
            }
        ],
        "knowledge": [
            {
                "knowledge_id": "KP-HRV-RESERVE-001",
                "title": "HRV恢复储备：身体的能量银行",
                "summary": "心率变异性(HRV)是衡量自主神经系统健康的重要指标。SDNN和RMSSD反映身体的恢复能力，就像您的'能量银行余额'。当HRV偏低但主观感觉良好时，说明身体正在消耗储备能量，需要及时补充。长期忽视可能导致真正的身心疲惫。"
            },
            {
                "knowledge_id": "KP-HRV-RESERVE-002",
                "title": "隐性疲劳的识别与预防",
                "summary": "隐性疲劳是指身体已经进入疲劳状态，但大脑尚未察觉的现象。常见于高压工作者和长期处于'战斗模式'的人群。通过HRV监测可以早期发现，采取主动休息可以有效预防。"
            }
        ],
        "videos": [
            {
                "video_id": "VID-REST-DEMO-001",
                "title": "科学休息示范",
                "description": "15分钟主动休息引导，包含呼吸调节和身体放松技巧",
                "duration_seconds": 900,
                "path": "./assets/videos/rest_demo_01.mp4"
            }
        ],
        "products": [
            {
                "product_id": "PROD-HAITANG-WEARABLE-001",
                "name": "海棠心智穿戴设备",
                "relevance": "实时监测HRV恢复储备，提供个性化休息建议，帮助追踪隐性疲劳改善进展"
            }
        ],
        "warnings": [
            "避免忽视身体的早期预警信号",
            "减少连续高强度工作时长",
            "保证每晚7-8小时睡眠",
            "如HRV持续偏低超过2周，建议咨询专业人士"
        ]
    },

    InterventionStrategy.MICRO_HABIT: {
        "name": "微习惯启动策略",
        "description": "针对低动机状态，从最小的改变开始",
        "priority": "normal",
        "max_tasks": 2,
        "max_difficulty": 1,
        "tone": "encouraging_gentle",
        "script": {
            "opening": "改变不需要很大，从最小的一步开始就好。重要的是开始，而不是完美。",
            "motivation": "每一个小小的行动，都在为大的改变铺路。不要小看这些微小的开始。",
            "closing": "这周就专注于这一个小习惯。做到了，就是成功。"
        },
        "tasks": [
            {
                "content": "每天喝一杯温水",
                "description": "早起后第一件事，喝一杯温水唤醒身体",
                "difficulty": 1,
                "type": "micro_habit",
                "duration_minutes": 1,
                "frequency": "每日"
            },
            {
                "content": "每天站立伸展1分钟",
                "description": "找任何一个时间，站起来伸展一下身体",
                "difficulty": 1,
                "type": "micro_habit",
                "duration_minutes": 1,
                "frequency": "每日"
            }
        ],
        "knowledge": [
            {
                "knowledge_id": "KP-MICRO-001",
                "title": "微习惯的力量",
                "summary": "为什么小到不可能失败的习惯，是改变的最佳起点。"
            }
        ],
        "videos": [
            {
                "video_id": "VID-MICRO-001",
                "title": "1分钟站立伸展",
                "description": "简单的站立伸展动作，随时随地可以做",
                "duration_seconds": 60
            }
        ],
        "products": [],
        "warnings": []
    },

    InterventionStrategy.GRADUAL_BUILDUP: {
        "name": "渐进式提升策略",
        "description": "准备期用户，逐步增加行为难度和频率",
        "priority": "normal",
        "max_tasks": 3,
        "max_difficulty": 2,
        "tone": "encouraging_practical",
        "script": {
            "opening": "您已经有了改变的意愿，现在让我们一步一步把它变成现实。",
            "motivation": "每周进步一点点，一个月后回头看，您会惊讶于自己的变化。",
            "closing": "这周的目标是可以达到的。相信自己，行动起来。"
        },
        "tasks": [
            {
                "content": "每天步行15分钟",
                "description": "可以分成早晚两次，每次7-8分钟",
                "difficulty": 2,
                "type": "exercise",
                "duration_minutes": 15,
                "frequency": "每日"
            },
            {
                "content": "记录每日三餐",
                "description": "用手机拍照或简单记录，不需要计算热量",
                "difficulty": 1,
                "type": "nutrition",
                "duration_minutes": 3,
                "frequency": "每日"
            },
            {
                "content": "睡前10分钟放松",
                "description": "可以是阅读、听音乐或简单的拉伸",
                "difficulty": 2,
                "type": "mental",
                "duration_minutes": 10,
                "frequency": "每日"
            }
        ],
        "knowledge": [
            {
                "knowledge_id": "KP-GRADUAL-001",
                "title": "习惯养成的科学",
                "summary": "了解习惯形成的神经机制，掌握建立习惯的最佳方法。"
            }
        ],
        "videos": [
            {
                "video_id": "VID-GRADUAL-001",
                "title": "室内简易有氧操",
                "description": "15分钟居家有氧运动，不需要任何器械",
                "duration_seconds": 900
            }
        ],
        "products": [
            {
                "product_id": "PROD-GRADUAL-001",
                "name": "计步手环",
                "relevance": "追踪每日步数，提供运动反馈"
            }
        ],
        "warnings": []
    },

    InterventionStrategy.HABIT_STRENGTHENING: {
        "name": "习惯强化策略",
        "description": "行动期用户，强化已有习惯，建立系统",
        "priority": "normal",
        "max_tasks": 4,
        "max_difficulty": 3,
        "tone": "supportive_systematic",
        "script": {
            "opening": "您的改变已经在发生！现在让我们把这些好习惯变得更稳固。",
            "motivation": "坚持到现在的您已经证明了自己的能力。继续保持，让习惯成为自然。",
            "closing": "这周继续保持，同时注意识别可能的障碍并提前准备应对方案。"
        },
        "tasks": [
            {
                "content": "每天30分钟有氧运动",
                "description": "快走、慢跑、游泳或骑车，保持心率在中等强度",
                "difficulty": 3,
                "type": "exercise",
                "duration_minutes": 30,
                "frequency": "每日"
            },
            {
                "content": "每餐保证一份蔬菜",
                "description": "每餐至少有一份蔬菜，优先选择深绿色蔬菜",
                "difficulty": 2,
                "type": "nutrition",
                "duration_minutes": 0,
                "frequency": "每餐"
            },
            {
                "content": "固定作息时间",
                "description": "每天同一时间起床和睡觉，包括周末",
                "difficulty": 3,
                "type": "sleep",
                "duration_minutes": 0,
                "frequency": "每日"
            },
            {
                "content": "每周总结与计划",
                "description": "每周日花15分钟回顾本周进展，规划下周目标",
                "difficulty": 2,
                "type": "mental",
                "duration_minutes": 15,
                "frequency": "每周"
            }
        ],
        "knowledge": [
            {
                "knowledge_id": "KP-HABIT-001",
                "title": "习惯的环境设计",
                "summary": "如何设计支持性环境，让好习惯更容易坚持。"
            }
        ],
        "videos": [
            {
                "video_id": "VID-HABIT-001",
                "title": "HIIT间歇训练入门",
                "description": "20分钟高效燃脂训练，适合有一定基础的用户",
                "duration_seconds": 1200
            }
        ],
        "products": [
            {
                "product_id": "PROD-HABIT-001",
                "name": "智能体脂秤",
                "relevance": "追踪身体成分变化，提供数据反馈"
            }
        ],
        "warnings": []
    },

    InterventionStrategy.IDENTITY_CONSOLIDATION: {
        "name": "身份巩固策略",
        "description": "维持期用户，深化身份认同，成为榜样",
        "priority": "normal",
        "max_tasks": 5,
        "max_difficulty": 4,
        "tone": "empowering",
        "script": {
            "opening": "健康的生活方式已经成为您的一部分。现在是时候深化这个身份，并帮助他人。",
            "motivation": "您不再是在'养成习惯'，这些就是您。分享您的经验，可以帮助更多人。",
            "closing": "保持现在的状态，适度挑战自己，也考虑如何帮助身边的人。"
        },
        "tasks": [
            {
                "content": "保持当前运动习惯",
                "description": "继续每周至少150分钟中等强度运动",
                "difficulty": 3,
                "type": "exercise",
                "duration_minutes": 150,
                "frequency": "每周"
            },
            {
                "content": "尝试一项新的健康活动",
                "description": "学习瑜伽、游泳或其他您感兴趣的运动",
                "difficulty": 3,
                "type": "growth",
                "duration_minutes": 60,
                "frequency": "每周"
            },
            {
                "content": "分享健康心得",
                "description": "在群组或朋友圈分享您的健康经验和收获",
                "difficulty": 2,
                "type": "social",
                "duration_minutes": 10,
                "frequency": "每周"
            },
            {
                "content": "帮助一位朋友开始健康改变",
                "description": "分享您的经验，陪伴他们迈出第一步",
                "difficulty": 4,
                "type": "social",
                "duration_minutes": 30,
                "frequency": "每周"
            }
        ],
        "knowledge": [
            {
                "knowledge_id": "KP-IDENTITY-001",
                "title": "从习惯到身份",
                "summary": "如何让健康行为成为自我认同的一部分，实现持久改变。"
            }
        ],
        "videos": [
            {
                "video_id": "VID-IDENTITY-001",
                "title": "高级训练技巧",
                "description": "进阶训练方法，适合已有稳定基础的用户",
                "duration_seconds": 1800
            }
        ],
        "products": [],
        "warnings": []
    }
}


# ============ generate_prescription 函数 ============

def generate_prescription(
    physio_data: Dict[str, Any],
    psych_data: Dict[str, Any],
    history: Dict[str, Any] = None,
    rx_library_path: str = "models/rx_library.json"
) -> Dict[str, Any]:
    """
    生成行为处方

    输入: data_extractor.py 提取的生理和心理指标

    逻辑:
    1. 如果 SDNN < 30 且 疲劳度 > 70 → 机体耗竭模式 → 强制性修复处方
    2. 如果 精力 < 40 → 微习惯启动策略
    3. 其他情况根据改变阶段匹配处方库

    输出: Agent 可读的 JSON 结构

    Args:
        physio_data: 生理指标字典 (来自 data_extractor)
        psych_data: 心理指标字典 (来自 data_extractor)
        history: 用户历史数据 (可选)
        rx_library_path: 处方库路径

    Returns:
        Agent 可读的处方 JSON
    """
    history = history or {}

    # 1. 创建行为画像
    profile = BehavioralProfile.from_extraction_result(physio_data, psych_data, history)

    print("\n" + "="*60)
    print("【处方引擎】行为画像分析")
    print("="*60)
    print(f"  行为模式: {profile.behavior_mode.value}")
    print(f"  动机水平: {profile.motivation_level.value} (分数: {profile.motivation_score:.1f})")
    print(f"  改变阶段: {profile.change_stage.value}")
    print(f"  干预策略: {profile.intervention_strategy.value}")
    print(f"  风险等级: {profile.risk_level}")

    if profile.is_exhaustion_mode:
        print(f"  ⚠️ 检测到机体耗竭模式 (SDNN={profile.sdnn}, 疲劳度={profile.fatigue_index})")

    if profile.is_hidden_fatigue_mode:
        print(f"  ⚠️ 检测到隐性疲劳模式 (SDNN={profile.sdnn}, 压力指数={profile.stress_index})")

    if profile.is_low_motivation:
        print(f"  ⚠️ 检测到低动机状态 (精力={profile.energy_level})")

    # 2. 获取处方模板
    template = PRESCRIPTION_TEMPLATES.get(
        profile.intervention_strategy,
        PRESCRIPTION_TEMPLATES[InterventionStrategy.GRADUAL_BUILDUP]
    )

    # 3. 尝试从 rx_library 获取补充内容
    rx_content = _load_rx_library_content(rx_library_path, profile)

    # 4. 生成处方 ID
    now = datetime.now()
    prescription_id = f"RX-{now.strftime('%Y%m%d')}-{uuid.uuid4().hex[:6].upper()}"

    # 5. 构建 Agent 可读的 JSON 结构
    prescription = {
        "prescription_id": prescription_id,
        "generated_at": now.isoformat(),
        "valid_until": (now + timedelta(days=7)).isoformat(),
        "version": "2.0",

        # 行为画像摘要
        "behavioral_profile": profile.to_dict(),

        # 处方元信息
        "prescription_meta": {
            "name": template["name"],
            "description": template["description"],
            "priority": template["priority"],
            "intervention_strategy": profile.intervention_strategy.value,
            "tone": template["tone"],
            "max_tasks": template["max_tasks"],
            "max_difficulty": template["max_difficulty"]
        },

        # 对话脚本 (Agent 使用)
        "conversation_script": template["script"],

        # 处方内容
        "prescription_content": {
            "guidance": {
                "summary": template["script"]["opening"][:100],
                "opening": template["script"]["opening"],
                "motivation": template["script"]["motivation"],
                "closing": template["script"]["closing"],
                "warnings": template.get("warnings", []),
                "do": _get_do_list(profile),
                "dont": _get_dont_list(profile)
            },

            "tasks": _generate_tasks_with_ids(template["tasks"], template["max_tasks"]),

            "knowledge": rx_content.get("knowledge", template.get("knowledge", [])),

            "videos": rx_content.get("videos", template.get("videos", [])),

            "products": rx_content.get("products", template.get("products", []))
        },

        # 限幅信息
        "clamping_info": {
            "max_tasks_allowed": template["max_tasks"],
            "max_difficulty_allowed": template["max_difficulty"],
            "clamping_reason": _get_clamping_reason(profile)
        },

        # Agent 指令
        "agent_instructions": {
            "tone": template["tone"],
            "key_messages": [
                template["script"]["opening"],
                template["script"]["motivation"]
            ],
            "avoid": template.get("warnings", []),
            "follow_up_questions": _get_follow_up_questions(profile)
        },

        # Agent 指导意见 (基于动机强度和改变阶段的 Coach 话术)
        "agent_guidance": generate_agent_guidance(profile)
    }

    print("\n" + "="*60)
    print("【处方引擎】处方生成完成")
    print("="*60)
    print(f"  处方ID: {prescription_id}")
    print(f"  处方类型: {template['name']}")
    print(f"  任务数量: {len(prescription['prescription_content']['tasks'])}")
    print(f"  知识点: {len(prescription['prescription_content']['knowledge'])}")
    print(f"  视频: {len(prescription['prescription_content']['videos'])}")
    print(f"  产品: {len(prescription['prescription_content']['products'])}")

    return prescription


def _load_rx_library_content(rx_library_path: str, profile: BehavioralProfile) -> Dict[str, List]:
    """从 rx_library 加载补充内容"""
    result = {"knowledge": [], "videos": [], "products": []}

    try:
        path = Path(rx_library_path)
        if not path.exists():
            return result

        with open(path, 'r', encoding='utf-8') as f:
            rx_library = json.load(f)

        prescriptions = rx_library.get("prescriptions", [])

        # 根据行为模式匹配处方
        for rx in prescriptions:
            applicable = rx.get("applicable_patterns", [])
            if profile.behavior_mode.value in applicable:
                content = rx.get("content", {})

                # 获取内容
                if content.get("knowledge_points"):
                    result["knowledge"].extend(content["knowledge_points"][:2])
                if content.get("teaching_videos"):
                    result["videos"].extend(content["teaching_videos"][:2])
                if rx.get("product_mapping"):
                    result["products"].extend(rx["product_mapping"][:2])

                break  # 只取第一个匹配的处方

    except Exception as e:
        print(f"[警告] 加载处方库失败: {e}")

    return result


def _generate_tasks_with_ids(tasks: List[Dict], max_tasks: int) -> List[Dict]:
    """为任务生成唯一 ID"""
    result = []
    for i, task in enumerate(tasks[:max_tasks]):
        result.append({
            "task_id": f"T-{uuid.uuid4().hex[:6].upper()}",
            "content": task["content"],
            "description": task.get("description", ""),
            "difficulty": task["difficulty"],
            "type": task["type"],
            "duration_minutes": task["duration_minutes"],
            "frequency": task["frequency"],
            "priority": i + 1
        })
    return result


def _get_do_list(profile: BehavioralProfile) -> List[str]:
    """根据画像生成 DO 列表"""
    do_map = {
        InterventionStrategy.MANDATORY_RECOVERY: ["优先休息", "倾听身体", "减少活动"],
        InterventionStrategy.HIDDEN_FATIGUE_RECOVERY: ["监测HRV", "主动休息", "早期预防"],
        InterventionStrategy.MICRO_HABIT: ["从最小开始", "庆祝每一步", "保持耐心"],
        InterventionStrategy.GRADUAL_BUILDUP: ["设定小目标", "逐步增加", "记录进展"],
        InterventionStrategy.HABIT_STRENGTHENING: ["保持规律", "设计环境", "预防复发"],
        InterventionStrategy.IDENTITY_CONSOLIDATION: ["分享经验", "帮助他人", "持续学习"]
    }
    return do_map.get(profile.intervention_strategy, ["保持积极", "循序渐进"])


def _get_dont_list(profile: BehavioralProfile) -> List[str]:
    """根据画像生成 DONT 列表"""
    dont_map = {
        InterventionStrategy.MANDATORY_RECOVERY: ["高强度运动", "熬夜工作", "忽视疲劳信号"],
        InterventionStrategy.HIDDEN_FATIGUE_RECOVERY: ["忽视早期信号", "继续透支", "依赖咖啡因"],
        InterventionStrategy.MICRO_HABIT: ["目标过大", "急于求成", "自我批评"],
        InterventionStrategy.GRADUAL_BUILDUP: ["一次改太多", "忽视困难", "完美主义"],
        InterventionStrategy.HABIT_STRENGTHENING: ["松懈大意", "忽视障碍", "过度自信"],
        InterventionStrategy.IDENTITY_CONSOLIDATION: ["过度控制", "忽视变化", "制造依赖"]
    }
    return dont_map.get(profile.intervention_strategy, ["急于求成", "自我否定"])


def _get_clamping_reason(profile: BehavioralProfile) -> str:
    """获取限幅原因"""
    if profile.is_exhaustion_mode:
        return "机体耗竭模式，强制限制任务强度以保护身体"
    elif profile.is_hidden_fatigue_mode:
        return "隐性疲劳模式，HRV恢复储备偏低，限制任务强度以预防耗竭"
    elif profile.is_low_motivation:
        return "低动机状态，采用微习惯策略降低门槛"
    elif profile.change_stage == ChangeStage.CONTEMPLATION:
        return "意向期，降低任务难度以建立信心"
    elif profile.change_stage == ChangeStage.PREPARATION:
        return "准备期，适度任务难度以逐步建立习惯"
    else:
        return "正常限幅"


def _get_follow_up_questions(profile: BehavioralProfile) -> List[str]:
    """生成后续跟进问题"""
    questions = {
        InterventionStrategy.MANDATORY_RECOVERY: [
            "您昨晚睡得怎么样？",
            "今天身体感觉如何？",
            "有什么让您特别疲惫的事情吗？"
        ],
        InterventionStrategy.HIDDEN_FATIGUE_RECOVERY: [
            "今天的HRV数据看了吗？",
            "午间休息感觉怎么样？",
            "有注意到身体的细微疲劳信号吗？"
        ],
        InterventionStrategy.MICRO_HABIT: [
            "今天的小任务完成了吗？",
            "有什么让您感到困难的地方吗？",
            "完成后您的感受是什么？"
        ],
        InterventionStrategy.GRADUAL_BUILDUP: [
            "这周的目标进展如何？",
            "遇到了什么挑战吗？",
            "下周想尝试什么新的改变？"
        ],
        InterventionStrategy.HABIT_STRENGTHENING: [
            "习惯保持得怎么样？",
            "有没有遇到想要放弃的时候？",
            "环境设计有帮助吗？"
        ],
        InterventionStrategy.IDENTITY_CONSOLIDATION: [
            "这周有分享您的经验吗？",
            "帮助他人的感觉如何？",
            "有什么新的健康目标吗？"
        ]
    }
    return questions.get(profile.intervention_strategy, ["今天感觉怎么样？"])


# ============ Agent 指导意见生成 ============

def generate_agent_guidance(profile: BehavioralProfile) -> Dict[str, Any]:
    """
    生成 Agent 指导意见

    基于动机强度和改变阶段，生成针对性的 Coach 话术

    Args:
        profile: BehavioralProfile 实例

    Returns:
        Agent 可读的指导意见 JSON
    """
    # 动机强度话术库
    motivation_scripts = {
        MotivationLevel.DEPLETED: {
            "tone": "温和关怀",
            "approach": "以陪伴为主，不给压力",
            "opening_phrases": [
                "我理解现在可能感觉有点累...",
                "没关系，我们慢慢来...",
                "今天不需要做很多，只要一小步就好..."
            ],
            "encouragement": [
                "能够开始就已经很棒了",
                "休息也是进步的一部分",
                "照顾好自己是最重要的"
            ],
            "avoid": ["催促", "高期望", "比较"]
        },
        MotivationLevel.LOW: {
            "tone": "轻柔鼓励",
            "approach": "降低门槛，增加成功体验",
            "opening_phrases": [
                "我们从最简单的开始...",
                "只需要一分钟，试试看？",
                "做到一点点就是胜利..."
            ],
            "encouragement": [
                "每一小步都在积累",
                "坚持比完美更重要",
                "您正在建立新的习惯"
            ],
            "avoid": ["完美主义", "长期目标", "复杂任务"]
        },
        MotivationLevel.MODERATE: {
            "tone": "积极支持",
            "approach": "稳步推进，建立节奏",
            "opening_phrases": [
                "今天的目标很清晰...",
                "让我们继续保持这个节奏...",
                "您做得很好，继续前进..."
            ],
            "encouragement": [
                "进步是可见的",
                "您的努力正在产生效果",
                "保持这个状态"
            ],
            "avoid": ["过度加量", "忽视困难"]
        },
        MotivationLevel.HIGH: {
            "tone": "激励挑战",
            "approach": "适度提升挑战，巩固身份",
            "opening_phrases": [
                "准备好迎接新挑战了吗？",
                "您的状态很好，我们可以...",
                "是时候尝试一些新的了..."
            ],
            "encouragement": [
                "您已经证明了自己的能力",
                "健康已经成为您的一部分",
                "继续探索您的潜力"
            ],
            "avoid": ["过度自信", "忽视休息"]
        }
    }

    # 改变阶段话术库
    stage_scripts = {
        ChangeStage.PRECONTEMPLATION: {
            "goal": "唤起意识，不施压",
            "key_message": "了解自己是改变的第一步",
            "conversation_style": "探索式提问",
            "sample_questions": [
                "您有没有注意到最近身体的一些变化？",
                "如果有机会改善一个方面，您会选择什么？"
            ]
        },
        ChangeStage.CONTEMPLATION: {
            "goal": "强化动机，解决矛盾",
            "key_message": "改变是可能的，而且值得",
            "conversation_style": "动机访谈",
            "sample_questions": [
                "是什么让您开始考虑改变？",
                "如果改变成功，您的生活会有什么不同？"
            ]
        },
        ChangeStage.PREPARATION: {
            "goal": "制定计划，建立信心",
            "key_message": "准备充分，成功率更高",
            "conversation_style": "计划导向",
            "sample_questions": [
                "您觉得从哪里开始最容易？",
                "有什么可能的障碍，我们可以提前准备？"
            ]
        },
        ChangeStage.ACTION: {
            "goal": "支持执行，预防复发",
            "key_message": "坚持下去，习惯正在形成",
            "conversation_style": "支持性反馈",
            "sample_questions": [
                "这周执行得怎么样？",
                "遇到困难时您是怎么应对的？"
            ]
        },
        ChangeStage.MAINTENANCE: {
            "goal": "巩固成果，深化身份",
            "key_message": "这就是您，健康已经是您的一部分",
            "conversation_style": "身份强化",
            "sample_questions": [
                "回顾这段旅程，您最大的收获是什么？",
                "您愿意分享您的经验帮助他人吗？"
            ]
        }
    }

    # 获取当前状态的话术
    motivation_script = motivation_scripts.get(
        profile.motivation_level,
        motivation_scripts[MotivationLevel.MODERATE]
    )
    stage_script = stage_scripts.get(
        profile.change_stage,
        stage_scripts[ChangeStage.CONTEMPLATION]
    )

    # 生成综合 Coach 话术
    coach_script = _generate_coach_script(profile, motivation_script, stage_script)

    return {
        "motivation_guidance": {
            "level": profile.motivation_level.value,
            "score": round(profile.motivation_score, 1),
            "tone": motivation_script["tone"],
            "approach": motivation_script["approach"],
            "opening_phrases": motivation_script["opening_phrases"],
            "encouragement": motivation_script["encouragement"],
            "avoid": motivation_script["avoid"]
        },
        "stage_guidance": {
            "stage": profile.change_stage.value,
            "spi_coefficient": profile.spi_coefficient,
            "goal": stage_script["goal"],
            "key_message": stage_script["key_message"],
            "conversation_style": stage_script["conversation_style"],
            "sample_questions": stage_script["sample_questions"]
        },
        "coach_script": coach_script,
        "session_flow": _generate_session_flow(profile)
    }


def _generate_coach_script(
    profile: BehavioralProfile,
    motivation_script: Dict,
    stage_script: Dict
) -> Dict[str, str]:
    """生成综合 Coach 话术脚本"""

    # 根据隐性疲劳模式特别处理
    if profile.is_hidden_fatigue_mode:
        return {
            "opening": f"{motivation_script['opening_phrases'][0]} 您的身体数据提示有一些隐性疲劳，虽然您可能感觉还好。让我们一起关注一下。",
            "awareness": "隐性疲劳就像手机电量显示还有30%，但其实快没电了。HRV数据帮助我们看到真实的'电量'。",
            "action_prompt": "这周我们只需要做两件事：监测HRV和午间休息。简单但很重要。",
            "encouragement": motivation_script["encouragement"][0],
            "closing": "记住，识别问题就是解决问题的一半。您已经在正确的路上了。"
        }

    # 根据机体耗竭模式特别处理
    if profile.is_exhaustion_mode:
        return {
            "opening": "我看到您的身体正在发出强烈的信号。现在最重要的不是做更多，而是让身体真正休息。",
            "awareness": "您的HRV和疲劳指数都在提示需要紧急恢复。这不是懒惰，是身体的智慧。",
            "action_prompt": "这周只有一个任务：好好休息。工作可以等，健康不能等。",
            "encouragement": "照顾好自己是最勇敢的选择。",
            "closing": "身体恢复后，我们再一起规划下一步。现在，请优先休息。"
        }

    # 通用 Coach 话术生成
    opening = f"{motivation_script['opening_phrases'][0]} {stage_script['key_message']}"

    action_prompts = {
        ChangeStage.PRECONTEMPLATION: "让我们先了解一下您的现状，不需要做任何改变。",
        ChangeStage.CONTEMPLATION: "想想看，如果开始改变，您希望从哪里开始？",
        ChangeStage.PREPARATION: "我们已经准备好了计划，这周就从第一个小目标开始。",
        ChangeStage.ACTION: "继续保持您的节奏，遇到困难随时告诉我。",
        ChangeStage.MAINTENANCE: "您做得很好！现在让我们想想怎么把这些好习惯传递给身边的人。"
    }

    return {
        "opening": opening,
        "awareness": f"您当前处于{_get_stage_chinese(profile.change_stage)}，{stage_script['goal']}是这个阶段的重点。",
        "action_prompt": action_prompts.get(profile.change_stage, "让我们一起努力。"),
        "encouragement": motivation_script["encouragement"][0],
        "closing": f"记住：{stage_script['key_message']} 下次见面时，期待听到您的进展！"
    }


def _get_stage_chinese(stage: ChangeStage) -> str:
    """获取改变阶段的中文名称"""
    stage_names = {
        ChangeStage.PRECONTEMPLATION: "前意向期",
        ChangeStage.CONTEMPLATION: "意向期",
        ChangeStage.PREPARATION: "准备期",
        ChangeStage.ACTION: "行动期",
        ChangeStage.MAINTENANCE: "维持期"
    }
    return stage_names.get(stage, "过渡期")


def _generate_session_flow(profile: BehavioralProfile) -> List[Dict[str, str]]:
    """生成会话流程建议"""
    base_flow = [
        {"step": "1", "action": "问候与状态确认", "duration": "1分钟"},
        {"step": "2", "action": "回顾上周进展", "duration": "2分钟"},
        {"step": "3", "action": "本周重点讲解", "duration": "3分钟"},
        {"step": "4", "action": "任务布置与答疑", "duration": "3分钟"},
        {"step": "5", "action": "鼓励与下周预告", "duration": "1分钟"}
    ]

    # 根据动机水平调整
    if profile.motivation_level == MotivationLevel.DEPLETED:
        base_flow[2]["action"] = "简短分享一个小知识点"
        base_flow[2]["duration"] = "2分钟"
        base_flow[3]["action"] = "布置一个微小任务"

    # 根据隐性疲劳调整
    if profile.is_hidden_fatigue_mode:
        base_flow.insert(2, {
            "step": "2.5",
            "action": "HRV数据解读与隐性疲劳科普",
            "duration": "2分钟"
        })

    return base_flow


# ============ 便捷函数 ============

def quick_prescription(
    sdnn: float = 50.0,
    fatigue: float = 50.0,
    energy: float = 50.0,
    mood: float = 50.0,
    anxiety: float = 30.0
) -> Dict[str, Any]:
    """
    快速生成处方 (用于测试)

    示例:
    >>> rx = quick_prescription(sdnn=25, fatigue=80)  # 触发机体耗竭模式
    >>> rx = quick_prescription(energy=30)            # 触发微习惯策略
    """
    physio = {"SDNN": sdnn}
    psych = {
        "fatigue_index": fatigue,
        "energy": energy,
        "mood": mood,
        "anxiety_score": anxiety
    }
    return generate_prescription(physio, psych)


# ============ Word 文档导出 ============

def export_to_word(prescription: Dict[str, Any], output_path: str) -> bool:
    """
    将处方 JSON 导出为专业 Word 文档报告

    Args:
        prescription: 处方 JSON 数据
        output_path: 输出文件路径

    Returns:
        成功返回 True，失败返回 False
    """
    if not DOCX_AVAILABLE:
        print("[错误] 需要安装 python-docx: pip install python-docx")
        return False

    try:
        doc = Document()

        # 设置文档默认字体
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # ========== 封面/标题 ==========
        _add_report_header(doc, prescription)

        # ========== 行为画像摘要 ==========
        _add_behavioral_profile_section(doc, prescription)

        # ========== 处方信息 ==========
        _add_prescription_meta_section(doc, prescription)

        # ========== 核心任务 ==========
        _add_tasks_section(doc, prescription)

        # ========== Agent 指导意见 ==========
        _add_agent_guidance_section(doc, prescription)

        # ========== 知识参考 ==========
        _add_knowledge_section(doc, prescription)

        # ========== 视频资源 ==========
        _add_videos_section(doc, prescription)

        # ========== 产品推荐 ==========
        _add_products_section(doc, prescription)

        # ========== 注意事项 ==========
        _add_warnings_section(doc, prescription)

        # ========== 页脚 ==========
        _add_report_footer(doc, prescription)

        # 保存文档
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

        return True

    except Exception as e:
        print(f"[错误] Word 文档生成失败: {e}")
        return False


def _add_report_header(doc: 'Document', prescription: Dict):
    """添加报告标题头"""
    # 主标题
    title = doc.add_heading('行为健康处方报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 副标题 - 处方类型
    meta = prescription.get('prescription_meta', {})
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(meta.get('name', '个性化行为处方'))
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    # 处方编号和日期
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rx_id = prescription.get('prescription_id', 'N/A')
    generated_at = prescription.get('generated_at', '')[:10]
    valid_until = prescription.get('valid_until', '')[:10]
    info_para.add_run(f"处方编号: {rx_id}").font.size = Pt(10)
    info_para.add_run(f"  |  生成日期: {generated_at}").font.size = Pt(10)
    info_para.add_run(f"  |  有效期至: {valid_until}").font.size = Pt(10)

    doc.add_paragraph()  # 空行


def _add_behavioral_profile_section(doc: 'Document', prescription: Dict):
    """添加行为画像摘要部分"""
    doc.add_heading('一、行为画像分析', level=1)

    profile = prescription.get('behavioral_profile', {})
    analysis = profile.get('analysis', {})
    scores = profile.get('scores', {})
    flags = profile.get('flags', {})
    raw = profile.get('raw_metrics', {})

    # 创建核心指标表格
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    header_cells = table.rows[0].cells
    headers = ['指标类别', '当前状态', '评分', '说明']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 行为模式
    row = table.add_row().cells
    row[0].text = '行为模式'
    behavior_mode = analysis.get('behavior_mode', 'balanced')
    mode_names = {
        'exhaustion': '机体耗竭',
        'hidden_fatigue': '隐性疲劳',
        'overcompensation': '过度补偿',
        'stress_avoidance': '应激逃避',
        'somatization': '躯体化',
        'emotional_dysregulation': '情绪失调',
        'balanced': '平衡'
    }
    row[1].text = mode_names.get(behavior_mode, behavior_mode)
    row[2].text = '-'
    mode_desc = {
        'exhaustion': '身体严重透支，需紧急恢复',
        'hidden_fatigue': 'HRV低但主观感觉良好，需早期干预',
        'balanced': '状态良好，可正常推进'
    }
    row[3].text = mode_desc.get(behavior_mode, '需关注并调整')

    # 动机水平
    row = table.add_row().cells
    row[0].text = '动机水平'
    motivation = analysis.get('motivation_level', 'moderate')
    level_names = {'depleted': '耗竭', 'low': '低', 'moderate': '中等', 'high': '高'}
    row[1].text = level_names.get(motivation, motivation)
    row[2].text = str(scores.get('motivation_score', '-'))
    row[3].text = f"精力-心情匹配度: {scores.get('energy_mood_match', 0):.0%}"

    # 改变阶段
    row = table.add_row().cells
    row[0].text = '改变阶段'
    stage = analysis.get('change_stage', 'contemplation')
    stage_names = {
        'precontemplation': '前意向期',
        'contemplation': '意向期',
        'preparation': '准备期',
        'action': '行动期',
        'maintenance': '维持期'
    }
    row[1].text = stage_names.get(stage, stage)
    row[2].text = f"SPI系数: {scores.get('spi_coefficient', 0.7)}"
    row[3].text = f"干预策略: {analysis.get('intervention_strategy', '-')}"

    # 风险等级
    row = table.add_row().cells
    row[0].text = '风险等级'
    risk = analysis.get('risk_level', 'medium')
    risk_names = {'low': '低风险', 'medium': '中等风险', 'high': '高风险'}
    row[1].text = risk_names.get(risk, risk)
    row[2].text = '-'
    risk_desc = {'low': '状态稳定', 'medium': '需定期关注', 'high': '建议专业评估'}
    row[3].text = risk_desc.get(risk, '')

    doc.add_paragraph()

    # 关键标志提示
    if flags.get('is_hidden_fatigue_mode'):
        warning = doc.add_paragraph()
        run = warning.add_run('⚠️ 检测到隐性疲劳模式: ')
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)
        warning.add_run(f"HRV恢复储备偏低 (SDNN={raw.get('sdnn', 'N/A')}ms)，但主观压力感知正常。身体已在消耗储备能量，需及时调整。")

    if flags.get('is_exhaustion_mode'):
        warning = doc.add_paragraph()
        run = warning.add_run('🚨 机体耗竭警报: ')
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        warning.add_run("身体发出强烈恢复信号，请立即优先休息，暂停高强度活动。")

    doc.add_paragraph()


def _add_prescription_meta_section(doc: 'Document', prescription: Dict):
    """添加处方元信息部分"""
    doc.add_heading('二、处方概述', level=1)

    meta = prescription.get('prescription_meta', {})
    content = prescription.get('prescription_content', {})
    guidance = content.get('guidance', {})

    # 处方描述
    desc_para = doc.add_paragraph()
    desc_para.add_run('处方类型: ').font.bold = True
    desc_para.add_run(meta.get('name', '-'))

    desc_para = doc.add_paragraph()
    desc_para.add_run('适用说明: ').font.bold = True
    desc_para.add_run(meta.get('description', '-'))

    desc_para = doc.add_paragraph()
    desc_para.add_run('优先级: ').font.bold = True
    priority_names = {'urgent': '紧急', 'high': '高', 'normal': '普通'}
    desc_para.add_run(priority_names.get(meta.get('priority', 'normal'), '普通'))

    # 核心指导
    doc.add_heading('核心指导', level=2)

    if guidance.get('opening'):
        para = doc.add_paragraph()
        run = para.add_run('💡 ')
        para.add_run(guidance['opening'])

    if guidance.get('motivation'):
        para = doc.add_paragraph()
        run = para.add_run('💪 ')
        para.add_run(guidance['motivation'])

    # DO / DON'T 列表
    if guidance.get('do') or guidance.get('dont'):
        doc.add_paragraph()

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        header = table.rows[0].cells
        header[0].text = '✅ 建议做'
        header[1].text = '❌ 避免做'
        for cell in header:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        row = table.add_row().cells
        row[0].text = '\n'.join([f"• {item}" for item in guidance.get('do', [])])
        row[1].text = '\n'.join([f"• {item}" for item in guidance.get('dont', [])])

    doc.add_paragraph()


def _add_tasks_section(doc: 'Document', prescription: Dict):
    """添加任务列表部分"""
    doc.add_heading('三、本周任务', level=1)

    content = prescription.get('prescription_content', {})
    tasks = content.get('tasks', [])
    clamping = prescription.get('clamping_info', {})

    # 限幅说明
    if clamping.get('clamping_reason'):
        para = doc.add_paragraph()
        run = para.add_run('📋 任务限幅: ')
        run.font.bold = True
        para.add_run(clamping['clamping_reason'])
        para.add_run(f" (最大任务数: {clamping.get('max_tasks_allowed', 3)}, 最大难度: {clamping.get('max_difficulty_allowed', 3)})")

    doc.add_paragraph()

    if not tasks:
        doc.add_paragraph("暂无具体任务安排")
        return

    # 任务表格
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    header = table.rows[0].cells
    headers = ['优先级', '任务内容', '频率', '时长', '难度']
    for i, h in enumerate(headers):
        header[i].text = h
        header[i].paragraphs[0].runs[0].font.bold = True
        header[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for task in tasks:
        row = table.add_row().cells
        row[0].text = str(task.get('priority', '-'))
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[1].text = task.get('content', '-')
        row[2].text = task.get('frequency', '-')
        duration = task.get('duration_minutes', 0)
        row[3].text = f"{duration}分钟" if duration > 0 else '-'
        row[4].text = '⭐' * task.get('difficulty', 1)

    doc.add_paragraph()

    # 任务详细说明
    for i, task in enumerate(tasks, 1):
        if task.get('description'):
            para = doc.add_paragraph()
            para.add_run(f"任务{i}说明: ").font.bold = True
            para.add_run(task['description'])

    doc.add_paragraph()


def _add_agent_guidance_section(doc: 'Document', prescription: Dict):
    """添加 Agent 指导意见部分"""
    doc.add_heading('四、Coach 指导话术', level=1)

    agent_guidance = prescription.get('agent_guidance', {})
    coach_script = agent_guidance.get('coach_script', {})
    motivation_guidance = agent_guidance.get('motivation_guidance', {})
    stage_guidance = agent_guidance.get('stage_guidance', {})

    if not coach_script:
        doc.add_paragraph("暂无 Coach 指导话术")
        return

    # Coach 脚本流程
    script_items = [
        ('开场白', coach_script.get('opening', '')),
        ('觉察引导', coach_script.get('awareness', '')),
        ('行动提示', coach_script.get('action_prompt', '')),
        ('鼓励语', coach_script.get('encouragement', '')),
        ('结束语', coach_script.get('closing', ''))
    ]

    for label, content in script_items:
        if content:
            para = doc.add_paragraph()
            run = para.add_run(f"【{label}】")
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
            doc.add_paragraph(content)

    doc.add_paragraph()

    # 沟通风格指南
    doc.add_heading('沟通风格指南', level=2)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    header = table.rows[0].cells
    header[0].text = '维度'
    header[1].text = '建议'
    for cell in header:
        cell.paragraphs[0].runs[0].font.bold = True

    rows_data = [
        ('沟通语调', motivation_guidance.get('tone', '-')),
        ('沟通方式', motivation_guidance.get('approach', '-')),
        ('对话风格', stage_guidance.get('conversation_style', '-')),
        ('阶段目标', stage_guidance.get('goal', '-')),
        ('核心信息', stage_guidance.get('key_message', '-'))
    ]

    for label, value in rows_data:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value

    # 避免事项
    if motivation_guidance.get('avoid'):
        doc.add_paragraph()
        para = doc.add_paragraph()
        para.add_run('⚠️ 沟通中避免: ').font.bold = True
        para.add_run('、'.join(motivation_guidance['avoid']))

    doc.add_paragraph()


def _add_knowledge_section(doc: 'Document', prescription: Dict):
    """添加知识参考部分"""
    doc.add_heading('五、知识科普', level=1)

    content = prescription.get('prescription_content', {})
    knowledge = content.get('knowledge', [])

    if not knowledge:
        doc.add_paragraph("暂无相关知识推荐")
        return

    for i, item in enumerate(knowledge, 1):
        # 知识点标题
        title = item.get('title', item.get('knowledge_id', f'知识点{i}'))
        heading = doc.add_heading(f'{i}. {title}', level=2)

        # 知识点内容
        summary = item.get('summary', '')
        if summary:
            doc.add_paragraph(summary)

        # 知识点ID
        if item.get('knowledge_id'):
            para = doc.add_paragraph()
            para.add_run(f"参考编号: {item['knowledge_id']}").font.size = Pt(9)
            para.add_run().font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()


def _add_videos_section(doc: 'Document', prescription: Dict):
    """添加视频资源部分"""
    doc.add_heading('六、教学视频', level=1)

    content = prescription.get('prescription_content', {})
    videos = content.get('videos', [])

    if not videos:
        doc.add_paragraph("暂无相关视频推荐")
        return

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    header = table.rows[0].cells
    headers = ['视频名称', '描述', '时长', '路径']
    for i, h in enumerate(headers):
        header[i].text = h
        header[i].paragraphs[0].runs[0].font.bold = True

    for video in videos:
        row = table.add_row().cells
        row[0].text = video.get('title', '-')
        row[1].text = video.get('description', '-')
        duration = video.get('duration_seconds', 0)
        minutes = duration // 60
        seconds = duration % 60
        row[2].text = f"{minutes}分{seconds}秒" if duration > 0 else '-'
        row[3].text = video.get('path', video.get('video_id', '-'))

    doc.add_paragraph()


def _add_products_section(doc: 'Document', prescription: Dict):
    """添加产品推荐部分"""
    doc.add_heading('七、推荐产品', level=1)

    content = prescription.get('prescription_content', {})
    products = content.get('products', [])

    if not products:
        doc.add_paragraph("暂无相关产品推荐")
        return

    for product in products:
        para = doc.add_paragraph()
        name = product.get('name', product.get('product_id', '未知产品'))
        run = para.add_run(f"🛒 {name}")
        run.font.bold = True

        if product.get('relevance'):
            doc.add_paragraph(f"适用说明: {product['relevance']}")

        if product.get('product_id'):
            para = doc.add_paragraph()
            para.add_run(f"产品编号: {product['product_id']}").font.size = Pt(9)

    doc.add_paragraph()


def _add_warnings_section(doc: 'Document', prescription: Dict):
    """添加注意事项部分"""
    content = prescription.get('prescription_content', {})
    guidance = content.get('guidance', {})
    warnings = guidance.get('warnings', [])

    if not warnings:
        return

    doc.add_heading('八、注意事项', level=1)

    for warning in warnings:
        para = doc.add_paragraph()
        run = para.add_run('⚠️ ')
        para.add_run(warning)

    doc.add_paragraph()


def _add_report_footer(doc: 'Document', prescription: Dict):
    """添加报告页脚"""
    doc.add_paragraph()
    doc.add_paragraph('─' * 50)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('本报告由 行健行为教练 行为处方引擎 自动生成').font.size = Pt(9)

    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer2.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run('本报告仅供参考，不构成医疗诊断建议。如有健康问题，请咨询专业医疗人员。')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ============ 主函数入口 ============

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="行为处方引擎")
    parser.add_argument("--test", choices=["exhaustion", "hidden_fatigue", "low_energy", "normal", "action"],
                        default="normal", help="测试场景")
    parser.add_argument("--output", default=None, help="输出 JSON 文件路径")
    parser.add_argument("--word", default=None, help="输出 Word 文档路径 (如: report.docx)")
    parser.add_argument("--word-only", action="store_true", help="仅生成 Word 文档，不输出 JSON")

    args = parser.parse_args()

    # 测试场景
    test_scenarios = {
        "exhaustion": {
            "physio": {"SDNN": 25, "RMSSD": 15, "heart_rate": 85},
            "psych": {"fatigue_index": 80, "energy": 30, "mood": 40, "anxiety_score": 60}
        },
        "hidden_fatigue": {
            # 隐性疲劳场景: HRV恢复能力低 (SDNN=42) 但主观压力低 (stress_index=41)
            "physio": {"SDNN": 42, "RMSSD": 25, "heart_rate": 72},
            "psych": {"fatigue_index": 45, "energy": 55, "mood": 60, "anxiety_score": 35, "stress_index": 41}
        },
        "low_energy": {
            "physio": {"SDNN": 45, "RMSSD": 25, "heart_rate": 72},
            "psych": {"fatigue_index": 60, "energy": 35, "mood": 50, "anxiety_score": 40}
        },
        "normal": {
            "physio": {"SDNN": 55, "RMSSD": 35, "heart_rate": 70},
            "psych": {"fatigue_index": 45, "energy": 60, "mood": 65, "anxiety_score": 35}
        },
        "action": {
            "physio": {"SDNN": 70, "RMSSD": 45, "heart_rate": 68},
            "psych": {"fatigue_index": 30, "energy": 75, "mood": 80, "anxiety_score": 25},
            "history": {"assessment_count": 8, "consecutive_improvement_weeks": 4, "task_completion_rate": 0.85}
        }
    }

    scenario = test_scenarios[args.test]

    print(f"\n{'='*70}")
    print(f"行为处方引擎测试 - 场景: {args.test}")
    print(f"{'='*70}")

    prescription = generate_prescription(
        scenario["physio"],
        scenario["psych"],
        scenario.get("history")
    )

    # 输出处方
    # 1. Word 文档输出
    if args.word:
        word_path = args.word
        if not word_path.endswith('.docx'):
            word_path += '.docx'
        print(f"\n正在生成 Word 报告...")
        if export_to_word(prescription, word_path):
            print(f"✅ Word 报告已保存: {word_path}")
        else:
            print(f"❌ Word 报告生成失败")

    # 2. JSON 输出
    if args.output:
        output_path = Path(args.output)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(prescription, f, ensure_ascii=False, indent=2)
        print(f"✅ JSON 处方已保存: {output_path}")

    # 3. 控制台预览 (如果没有指定 --word-only)
    if not args.word_only and not args.output:
        print("\n处方预览:")
        print(json.dumps(prescription, ensure_ascii=False, indent=2)[:3000])
        if len(json.dumps(prescription)) > 3000:
            print("\n... (输出已截断)")

    # 4. 自动生成 Word (如果只指定了 --word-only 但没指定 --word)
    if args.word_only and not args.word:
        auto_word_path = f"output/prescription_{args.test}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        print(f"\n正在生成 Word 报告...")
        if export_to_word(prescription, auto_word_path):
            print(f"✅ Word 报告已保存: {auto_word_path}")
        else:
            print(f"❌ Word 报告生成失败")
